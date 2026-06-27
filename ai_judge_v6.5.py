import os
import re
import shutil
import hashlib
import logging
import json
import datetime
import zipfile
import streamlit as st

# ==========================================
# LANGSMITH ТРЕЙСИНГ
# Активируется при вводе ключа в боковой панели.
# Отключить: оставить поле пустым.
# ==========================================
_ls_key = os.environ.get("LANGCHAIN_API_KEY", "")
if _ls_key:
    os.environ.setdefault("LANGCHAIN_TRACING_V2", "true")
    os.environ.setdefault("LANGCHAIN_PROJECT", "legal-rag-plenum")
    os.environ.setdefault("LANGCHAIN_ENDPOINT", "https://api.smith.langchain.com")
else:
    os.environ["LANGCHAIN_TRACING_V2"] = "false"
import torch
import chromadb
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader, PyPDFLoader, Docx2txtLoader
try:
    from langchain.retrievers import ParentDocumentRetriever
except ImportError:
    from langchain_classic.retrievers import ParentDocumentRetriever
try:
    from langchain_community.storage import LocalFileStore
except ImportError:
    try:
        from langchain.storage import LocalFileStore
    except ImportError:
        from langchain_classic.storage import LocalFileStore
try:
    from langchain.storage._lc_store import create_kv_docstore
except ImportError:
    from langchain_classic.storage._lc_store import create_kv_docstore

from typing import List
try:
    from langchain_community.retrievers import BM25Retriever
except ImportError:
    from langchain.retrievers import BM25Retriever

from langchain_core.documents import Document as LangChainDocument
from langchain_core.retrievers import BaseRetriever
from langchain_core.callbacks import CallbackManagerForRetrieverRun
from openai import OpenAI
import google.generativeai as genai
from google.generativeai.types import GenerationConfig
from io import BytesIO

try:
    from sentence_transformers import CrossEncoder
except ImportError:
    st.error("❌ Ошибка: Библиотека 'sentence-transformers' не установлена.")
    st.stop()

try:
    import rank_bm25
except ImportError:
    st.error("❌ Ошибка: Библиотека 'rank_bm25' не установлена. pip install rank_bm25")
    st.stop()

try:
    import docx
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    st.error("❌ Ошибка: Библиотека 'python-docx' не установлена.")
    st.stop()

import psutil
import toml
from pathlib import Path

# ==========================================
# LOGGING
# ==========================================
logging.basicConfig(
    filename="rag_errors.log",
    level=logging.WARNING,
    format="%(asctime)s | %(levelname)s | %(message)s",
    encoding="utf-8",
)

def log_low_recall(query: str, top_score: float):
    logging.warning(f"LOW_RECALL | score={top_score:.3f} | query={query[:120]}")

# ==========================================
# ENSEMBLE RETRIEVER
# ==========================================
class EnsembleRetriever(BaseRetriever):
    """Reciprocal Rank Fusion (Cormack et al. 2009): score = Σ w/(rank+60)."""
    retrievers: List[BaseRetriever]
    weights: List[float]
    rrf_k: int = 60

    def _get_relevant_documents(
        self, query: str, *, run_manager: CallbackManagerForRetrieverRun = None
    ) -> List[LangChainDocument]:
        from collections import defaultdict
        rrf_scores: dict = defaultdict(float)
        doc_map: dict = {}
        for retriever, weight in zip(self.retrievers, self.weights):
            try:
                docs = retriever.invoke(query)
                for rank, doc in enumerate(docs):
                    key = doc.page_content[:200]
                    rrf_scores[key] += weight / (rank + self.rrf_k)
                    if key not in doc_map:
                        doc_map[key] = doc
            except Exception as e:
                print(f"Ошибка ретривера: {e}")
                continue
        sorted_keys = sorted(rrf_scores, key=lambda k: rrf_scores[k], reverse=True)
        result = []
        for key in sorted_keys[:10]:
            doc = doc_map[key]
            if not hasattr(doc, "metadata"):
                doc.metadata = {}
            doc.metadata["ensemble_score"] = rrf_scores[key]
            result.append(doc)
        return result


if not hasattr(st, "rerun"):
    st.rerun = st.experimental_rerun

# ==========================================
# КОНСТАНТЫ
# ==========================================
# ── Обычный гибридный поиск (fallback) ─────────────────────────
CHUNK_SIZE    = 1000
CHUNK_OVERLAP = 200

# ── Parent Document Retriever ─────────────────────────────────────────
# Оптимальные размеры для юридических текстов (Пленумы ВС РФ):
#   child  ~850 симв. (~200 токенов)  — один полный пункт Пленума,
#                                        достаточно для точного векторного поиска
#   parent ~3200 симв. (~800 токенов) — 3–4 пункта, связный смысловой раздел,
#                                        передаётся в LLM для анализа
# Источник: best practices для legal RAG (child 400-2000 симв., parent 2000-8000 симв.)
CHILD_CHUNK_SIZE     = 850
CHILD_CHUNK_OVERLAP  = 100
PARENT_CHUNK_SIZE    = 3200
PARENT_CHUNK_OVERLAP = 200

# ── Веса (Этап 1 плана 24.02) ───────────────────────────────────────
BM25_WEIGHT   = 0.55
VECTOR_WEIGHT = 0.45

BATCH_SIZE          = 100
MIN_RELEVANCE_SCORE = 0.30

BASE_DIR        = os.getcwd()
RAG_DIR         = os.path.join(BASE_DIR, "rag_db")
PDR_DIR         = os.path.join(BASE_DIR, "rag_db_pdr")
PDR_DOCSTORE    = os.path.join(BASE_DIR, "rag_db_pdr_docs")
SOURCE_DOCS_DIR = os.path.join(BASE_DIR, "source_docs")

COLLECTION_NAME     = "plenum_vs_rf"
COLLECTION_NAME_PDR = "plenum_vs_rf_pdr"

os.makedirs(RAG_DIR,         exist_ok=True)
os.makedirs(PDR_DIR,         exist_ok=True)
os.makedirs(PDR_DOCSTORE,    exist_ok=True)
os.makedirs(SOURCE_DOCS_DIR, exist_ok=True)

# ==========================================
# КОНФИГУРАЦИЯ ПРОВАЙДЕРОВ LLM
# ==========================================
LLM_PROVIDERS = {
    "Gemini": {
        "type": "gemini",
        "models": [
            "gemini-2.5-pro",
            "gemini-2.5-flash",
            "gemini-2.0-pro-exp",
            "gemini-1.5-pro",
        ],
        "key_session": "gemini_api_key",
        "key_saved":   "gemini_saved",
        "label":       "Gemini API Key:",
        "help":        "Получить на https://aistudio.google.com/app/apikey",
    },
    "OpenRouter": {
        "type": "openrouter",
        "base_url": "https://openrouter.ai/api/v1",
        "models": [
            "google/gemini-2.5-pro",
            "google/gemini-2.5-flash",
            "anthropic/claude-3.7-sonnet",
            "meta-llama/llama-4-maverick",
            "deepseek/deepseek-chat-v3-0324",
            "mistralai/mistral-large",
        ],
        "key_session": "openrouter_api_key",
        "key_saved":   "openrouter_saved",
        "label":       "OpenRouter API Key:",
        "help":        "Получить на https://openrouter.ai/keys",
    },
}

# ==========================================
# РЕРАНКЕРЫ
# ==========================================
RERANKER_CONFIGS = {
    "BGE v2-m3 (быстрый)": {
        "type": "local_crossencoder",
        "model_path": "./models/bge-reranker-v2-m3",
        "max_length": 1024,
        "description": "Быстрый, для коротких материалов (до 3 страниц)",
        "ram_min_gb": 3,
    },
    "BGE v2.5-gemma (мощный)": {
        "type": "local_crossencoder",
        "model_path": "./models/bge-reranker-v2.5-gemma",
        "max_length": 8192,
        "description": "Мощный, для длинных материалов (до 20 страниц)",
        "ram_min_gb": 6,
    },
    "Cohere API (облачный)": {
        "type": "cohere_api",
        "model_name": "rerank-multilingual-v3.0",
        "max_length": 4096,
        "description": "Облачный реранкер для слабых ПК (нужен интернет)",
        "ram_min_gb": 2,
    },
}

# ==========================================
# ИЕРАРХИЯ ИСТОЧНИКОВ ПРАВА (только гражданские дела)
# ==========================================
DOCUMENT_TYPE_WEIGHTS = {
    "Пленум_ВС": 1.5,
    "Обзор_судебной_практики_ВС": 1.4,
    "Обзор_правовых_позиций_ВС": 1.35,
    "Определение_ВС": 1.2,
    "Кодекс": 1.1,
    "ФЗ": 1.0,
    "Другое": 0.9,
}


def detect_document_type(source: str, content: str) -> str:
    source_lower = source.lower()
    content_lower = content[:1000].lower()

    if "пленум" in source_lower or "постановление пленума" in content_lower:
        return "Пленум_ВС"
    if (
        ("обзор судебной практики" in content_lower or "обзор практики" in content_lower)
        and ("верховн" in content_lower or "вс" in content_lower)
        and "правовых позиций" not in content_lower[:500]
    ):
        return "Обзор_судебной_практики_ВС"
    if "обзор правовых позиций" in content_lower[:500] and (
        "верховн" in content_lower or "вс" in content_lower
    ):
        return "Обзор_правовых_позиций_ВС"
    if ("определение" in source_lower or "постановление" in source_lower) and (
        "верховн" in source_lower or "вс рф" in source_lower
    ):
        if "пленум" not in content_lower[:200]:
            return "Определение_ВС"
    if any(w in source_lower for w in ["гк рф", "гк_рф", "гпк рф", "гпк_рф", "кас рф", "кодекс"]):
        return "Кодекс"
    if re.search(r"федеральный закон|фз\s*№|фз\s+от", content_lower):
        return "ФЗ"
    return "Другое"


# ==========================================
# КОНФИГ СТРАНИЦЫ
# ==========================================
st.set_page_config(
    page_title="Анализ решений суда v6.5 — Пленумы ВС РФ",
    layout="wide",
    page_icon="⚖️",
)

# ==========================================
# УПРАВЛЕНИЕ API-КЛЮЧАМИ
# ==========================================
def load_api_keys_from_secrets():
    secrets_path = Path(".streamlit/secrets.toml")
    if not secrets_path.exists():
        return False
    try:
        secrets = toml.load(secrets_path)
        for prov_cfg in LLM_PROVIDERS.values():
            env_key = prov_cfg["key_session"].upper()
            if env_key in secrets:
                st.session_state[prov_cfg["key_session"]] = secrets[env_key]
                st.session_state[prov_cfg["key_saved"]] = True
        if "COHERE_API_KEY" in secrets:
            st.session_state["cohere_api_key"] = secrets["COHERE_API_KEY"]
            st.session_state["cohere_saved"] = True
        return True
    except Exception as e:
        st.error(f"Ошибка чтения secrets.toml: {e}")
        return False


def save_api_keys_to_secrets():
    secrets_path = Path(".streamlit/secrets.toml")
    secrets_path.parent.mkdir(exist_ok=True)
    try:
        existing = {}
        if secrets_path.exists():
            existing = toml.load(secrets_path)
        saved = []
        for prov_cfg in LLM_PROVIDERS.values():
            val = st.session_state.get(prov_cfg["key_session"], "")
            if val:
                existing[prov_cfg["key_session"].upper()] = val
                st.session_state[prov_cfg["key_saved"]] = True
                saved.append(prov_cfg["key_session"])
        ck = st.session_state.get("cohere_api_key", "")
        if ck:
            existing["COHERE_API_KEY"] = ck
            st.session_state["cohere_saved"] = True
        with open(secrets_path, "w", encoding="utf-8") as f:
            toml.dump(existing, f)
        return True
    except Exception as e:
        st.error(f"❌ Ошибка сохранения ключей: {e}")
        return False


def render_api_keys_section():
    st.subheader("🔑 API ключи")
    for prov_name, prov_cfg in LLM_PROVIDERS.items():
        val = st.text_input(
            prov_cfg["label"],
            value=st.session_state.get(prov_cfg["key_session"], ""),
            type="password",
            key=f"input_{prov_cfg['key_session']}",
            help=prov_cfg["help"],
        )
        st.session_state[prov_cfg["key_session"]] = val
        if val:
            st.caption("✅ Сохранён" if st.session_state.get(prov_cfg["key_saved"]) else "⚠️ Введён (не сохранён)")
        else:
            st.caption("ℹ️ Не введён")

    ck = st.text_input(
        "Cohere API (реранкер):",
        value=st.session_state.get("cohere_api_key", ""),
        type="password",
        key="cohere_input",
    )
    st.session_state["cohere_api_key"] = ck
    st.caption("✅ Сохранён" if (ck and st.session_state.get("cohere_saved")) else ("⚠️ Введён" if ck else "ℹ️ Не введён"))

    st.divider()
    st.markdown("**🔬 LangSmith (трейсинг — необязательно)**")
    ls_key = st.text_input(
        "LangSmith API Key:",
        value=st.session_state.get("langsmith_api_key", ""),
        type="password",
        help="Получить бесплатно: https://smith.langchain.com → Settings → API Keys",
        key="langsmith_input",
    )
    st.session_state["langsmith_api_key"] = ls_key
    if ls_key:
        os.environ["LANGCHAIN_TRACING_V2"] = "true"
        os.environ["LANGCHAIN_PROJECT"]     = "legal-rag-plenum"
        os.environ["LANGCHAIN_API_KEY"]     = ls_key
        st.caption("✅ LangSmith активен → трейсы пишутся в проект `legal-rag-plenum`")
    else:
        os.environ["LANGCHAIN_TRACING_V2"] = "false"
        st.caption("ℹ️ Не введён — трейсинг отключён")

    st.divider()
    col1, col2 = st.columns(2)
    with col1:
        if st.button("💾 Сохранить все", use_container_width=True):
            if save_api_keys_to_secrets():
                st.success("✅ Ключи сохранены")
                st.rerun()
    with col2:
        if st.button("🗑️ Удалить", use_container_width=True):
            p = Path(".streamlit/secrets.toml")
            if p.exists():
                p.unlink()
            for prov_cfg in LLM_PROVIDERS.values():
                st.session_state[prov_cfg["key_saved"]] = False
            st.session_state["cohere_saved"] = False
            st.rerun()


# ==========================================
# ИНИЦИАЛИЗАЦИЯ
# ==========================================
if "keys_loaded" not in st.session_state:
    load_api_keys_from_secrets()
    st.session_state["keys_loaded"] = True

# ==========================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ФАЙЛОВ
# ==========================================
def get_loader(file_path):
    ext = file_path.split(".")[-1].lower()
    loaders = {
        "txt":  lambda: TextLoader(file_path, encoding="utf-8"),
        "pdf":  lambda: PyPDFLoader(file_path),
        "docx": lambda: Docx2txtLoader(file_path),
    }
    lf = loaders.get(ext)
    return lf() if lf else None


@st.cache_data(show_spinner=False)
def read_file_content(uploaded_file) -> str:
    try:
        ext = uploaded_file.name.split(".")[-1].lower()
        if ext == "txt":
            return uploaded_file.read().decode("utf-8")
        elif ext == "docx":
            doc = docx.Document(uploaded_file)
            parts = []
            for el in doc.element.body:
                if isinstance(el, CT_P):
                    para = Paragraph(el, doc)
                    if para.text.strip():
                        parts.append(para.text)
                elif isinstance(el, CT_Tbl):
                    table = Table(el, doc)
                    parts.append("\n--- НАЧАЛО ТАБЛИЦЫ ---\n")
                    for row in table.rows:
                        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                        parts.append("| " + " | ".join(cells) + " |")
                    parts.append("--- КОНЕЦ ТАБЛИЦЫ ---\n")
            return "\n".join(parts)
        return f"[Формат {ext} не поддерживается]"
    except Exception as e:
        return f"[Ошибка чтения: {e}]"


# ==========================================
# УМНОЕ ИЗВЛЕЧЕНИЕ СНИПЕТА ИЗ РЕШЕНИЯ
# (вместо простого [:500])
# ==========================================
def extract_legal_snippet(text: str, max_chars: int = 1800) -> str:
    """
    Берёт начало + середину + конец текста, чтобы покрыть
    вводную часть, мотивировку и резолютивную часть решения.
    """
    if not text or len(text) <= max_chars:
        return text
    part = max_chars // 3
    mid_start = len(text) // 2 - part // 2
    mid_end   = mid_start + part
    return (
        text[:part]
        + "\n[...]\n"
        + text[mid_start:mid_end]
        + "\n[...]\n"
        + text[-part:]
    )


# ==========================================
# EMBEDDINGS & RERANKERS
# ==========================================
@st.cache_resource
def get_embeddings_model():
    device_type = "cuda" if torch.cuda.is_available() else "cpu"
    local_path = "./models/jina-embeddings-v3"
    repo_id = "jinaai/jina-embeddings-v3"

    if os.path.exists(local_path):
        try:
            from sentence_transformers import SentenceTransformer

            model = SentenceTransformer(local_path, device=device_type, trust_remote_code=True)

            class STEmbeddings:
                def __init__(self, model):
                    self.model = model

                def embed_documents(self, texts):
                    vecs = self.model.encode(texts, normalize_embeddings=True, convert_to_numpy=True)
                    return vecs.tolist()

                def embed_query(self, text):
                    vec = self.model.encode([text], normalize_embeddings=True, convert_to_numpy=True)[0]
                    return vec.tolist()

            st.success(f"✅ Jina v3 загружена локально из {local_path}")
            return STEmbeddings(model)
        except Exception as e:
            st.warning(f"⚠️ Ошибка загрузки из {local_path}: {e}")

    st.info("⏳ Jina v3 не найдена локально. Загружаю из HuggingFace (~560 MB)...")
    try:
        from sentence_transformers import SentenceTransformer

        model = SentenceTransformer(repo_id, device=device_type, trust_remote_code=True, cache_folder='.cache/huggingface')

        class STEmbeddings:
            def __init__(self, model):
                self.model = model

            def embed_documents(self, texts):
                vecs = self.model.encode(texts, normalize_embeddings=True, convert_to_numpy=True)
                return vecs.tolist()

            def embed_query(self, text):
                vec = self.model.encode([text], normalize_embeddings=True, convert_to_numpy=True)[0]
                return vec.tolist()

        st.success("✅ Jina v3 загружена")
        st.info(f"💡 Для ускорения в будущем скачайте локально: huggingface-cli download {repo_id} --local-dir {local_path}")
        return STEmbeddings(model)
    except Exception as e:
        st.error(f"❌ Не удалось загрузить {repo_id}: {e}")
        return None


@st.cache_resource(ttl=3600)
def get_reranker_bge_v2m3():
    device_type = "cuda" if torch.cuda.is_available() else "cpu"
    local_path = "./models/bge-reranker-v2-m3"
    repo_id = "BAAI/bge-reranker-v2-m3"
    if os.path.exists(local_path):
        try:
            return CrossEncoder(local_path, device=device_type, max_length=1024)
        except Exception as e:
            st.warning(f"⚠️ {e}")
    try:
        return CrossEncoder(repo_id, device=device_type, max_length=1024)
    except Exception as e:
        st.error(f"❌ {repo_id}: {e}")
        return None


@st.cache_resource(ttl=3600)
def get_reranker_bge_v25_gemma():
    candidate_paths = [
        "./models/bge-reranker-v2.5-gemma",
        "./models/models--BAAI--bge-reranker-v2.5-gemma2-lightweight",
        "./models/bge-reranker-v2.5-gemma2-lightweight",
    ]
    model_path = next((p for p in candidate_paths if os.path.exists(p)), None)
    if not model_path:
        return None
    try:
        device_type = "cuda" if torch.cuda.is_available() else "cpu"
        return CrossEncoder(model_path, device=device_type, max_length=8192)
    except Exception as e:
        st.error(f"❌ BGE v2.5-gemma: {e}")
        return None


def get_reranker_cohere():
    api_key = st.session_state.get("cohere_api_key", "")
    if not api_key:
        return None
    try:
        import cohere
        return cohere.Client(api_key=api_key)
    except Exception:
        return None


def get_reranker_model(reranker_choice: str):
    if reranker_choice == "BGE v2-m3 (быстрый)":
        return get_reranker_bge_v2m3()
    if reranker_choice == "BGE v2.5-gemma (мощный)":
        return get_reranker_bge_v25_gemma()
    if reranker_choice == "Cohere API (облачный)":
        return get_reranker_cohere()
    return None


def get_available_rerankers():
    available = ["BGE v2-m3 (быстрый)"]
    gemma_paths = [
        "./models/bge-reranker-v2.5-gemma",
        "./models/models--BAAI--bge-reranker-v2.5-gemma2-lightweight",
        "./models/bge-reranker-v2.5-gemma2-lightweight",
    ]
    if any(os.path.exists(p) for p in gemma_paths):
        if psutil.virtual_memory().available / (1024 ** 3) >= 5:
            available.append("BGE v2.5-gemma (мощный)")
    if st.session_state.get("cohere_api_key"):
        available.append("Cohere API (облачный)")
    return available


# ==========================================
# ВЕКТОРНАЯ БАЗА
# ==========================================
def load_vector_db():
    if not os.path.exists(RAG_DIR) or not os.listdir(RAG_DIR):
        return None
    try:
        return Chroma(
            persist_directory=RAG_DIR,
            embedding_function=get_embeddings_model(),
            collection_name=COLLECTION_NAME,
        )
    except Exception as e:
        st.error(f"Ошибка загрузки базы: {e}")
        return None


@st.cache_resource(show_spinner="Инициализация BM25...")
def get_bm25_retriever():
    vs = load_vector_db()
    if not vs:
        return None
    try:
        data = vs.get()
        texts, metas = data["documents"], data["metadatas"]
        if not texts:
            return None
        docs = [LangChainDocument(page_content=t, metadata=m) for t, m in zip(texts, metas)]
        bm25 = BM25Retriever.from_documents(docs)
        bm25.k = 10
        return bm25
    except Exception as e:
        print(f"Ошибка BM25: {e}")
        return None


def generate_doc_ids(documents):
    ids = []
    for i, doc in enumerate(documents):
        source = doc.metadata.get("source", "unknown")
        content_hash = hashlib.md5(doc.page_content.encode("utf-8")).hexdigest()
        unique_str = f"{source}||{i}||{content_hash}"
        ids.append(hashlib.md5(unique_str.encode("utf-8")).hexdigest())
    return ids



# ==========================================
# PARENT DOCUMENT RETRIEVER (PDR)
# child  ~850 симв. → индексируются в Chroma (точный поиск)
# parent ~3200 симв. → LocalFileStore, передаются в LLM
# ==========================================

@st.cache_resource(show_spinner="Инициализация Parent Document Retriever...")
def get_parent_document_retriever(k: int = 7):
    try:
        embeddings = get_embeddings_model()
        if embeddings is None:
            st.error("PDR init error: embeddings model is None")
            return None

        child_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHILD_CHUNK_SIZE,
            chunk_overlap=CHILD_CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", "; ", ", ", " "],
        )
        parent_splitter = RecursiveCharacterTextSplitter(
            chunk_size=PARENT_CHUNK_SIZE,
            chunk_overlap=PARENT_CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", "; ", ", ", " "],
        )
        vectorstore = Chroma(
            persist_directory=PDR_DIR,
            embedding_function=embeddings,
            collection_name=COLLECTION_NAME_PDR,
        )
        fs = LocalFileStore(PDR_DOCSTORE)
        docstore = create_kv_docstore(fs)
        retriever = ParentDocumentRetriever(
            vectorstore=vectorstore,
            docstore=docstore,
            child_splitter=child_splitter,
            parent_splitter=parent_splitter,
            search_kwargs={"k": k * 3},
        )
        return retriever
    except Exception as e:
        st.error(f"PDR init error: {e}")
        return None

def pdr_has_documents() -> bool:
    """True если PDR-коллекция непуста."""
    if not os.path.exists(PDR_DIR) or not os.listdir(PDR_DIR):
        return False
    try:
        vs = Chroma(
            persist_directory=PDR_DIR,
            embedding_function=get_embeddings_model(),
            collection_name=COLLECTION_NAME_PDR,
        )
        return len(vs.get()["ids"]) > 0
    except Exception:
        return False


def add_docs_to_pdr(documents: list, status_fn=None) -> int:
    """Добавляет документы в PDR. PDR сам нарезает child+parent и сохраняет."""
    retriever = get_parent_document_retriever()
    if retriever is None:
        if status_fn:
            status_fn("❌ PDR retriever не создался. Смотрите сообщение PDR init error выше.")
        return 0
    added = 0
    for i in range(0, len(documents), BATCH_SIZE):
        batch = documents[i : i + BATCH_SIZE]
        try:
            retriever.add_documents(batch)
            added += len(batch)
            if status_fn:
                status_fn(f"PDR: {added}/{len(documents)} документов...")
        except Exception as e:
            if status_fn:
                status_fn(f"⚠️ PDR батч {i}: {e}")
    return added

def get_pdr_docs(query: str, context_snippet: str, k: int = 7) -> list:
    """Ищет по child-чанкам → возвращает родительские документы (~3200 симв.)."""
    retriever = get_parent_document_retriever(k=k)
    if retriever is None or not pdr_has_documents():
        return []
    query_with_prefix = f"retrieval.query: {query}"
    search_text = f"{query_with_prefix} {context_snippet}" if context_snippet else query_with_prefix
    try:
        return retriever.invoke(search_text)
    except Exception as e:
        print(f"Ошибка PDR поиска: {e}")
        return []


def rebuild_pdr(status_fn=None, progress_fn=None, target_pdr_dir=None, target_docstore=None) -> str:
    """Полная пересборка PDR из source_docs, при необходимости в новые shadow-папки."""
    release_chroma_resources()
    pdr_dir = target_pdr_dir or make_shadow_dir(os.path.join(BASE_DIR, "rag_db_pdr"))
    docstore_dir = target_docstore or make_shadow_dir(os.path.join(BASE_DIR, "rag_db_pdr_docs"))
    os.makedirs(pdr_dir, exist_ok=True)
    os.makedirs(docstore_dir, exist_ok=True)

    old_pdr_dir = PDR_DIR
    old_docstore = PDR_DOCSTORE
    set_active_db_paths(pdr_dir=pdr_dir, pdr_docstore=docstore_dir)
    st.cache_resource.clear()

    try:
        files = [f for f in os.listdir(SOURCE_DOCS_DIR) if f.endswith((".txt", ".pdf", ".docx"))]
        if not files:
            return "❌ Папка source_docs пуста."

        documents = []
        total_files = len(files)
        for i, fname in enumerate(files):
            path = os.path.join(SOURCE_DOCS_DIR, fname)
            try:
                loader = get_loader(path)
                if loader:
                    loaded = loader.load()
                    for d in loaded:
                        d.metadata["indexed_at"] = datetime.datetime.now().isoformat()
                        d.metadata["doc_version"] = "v1"
                    documents.extend(loaded)
            except Exception as e:
                if status_fn:
                    status_fn(f"⚠️ {fname}: {e}")
            if progress_fn:
                progress_fn(0.05 + (i + 1) / max(total_files, 1) * 0.35)

        if not documents:
            return "❌ Не удалось загрузить файлы."

        if status_fn:
            status_fn(f"📦 Загружено {len(documents)} документов — нарезаю child+parent...")
        if progress_fn:
            progress_fn(0.45)

        n = add_docs_to_pdr(documents, status_fn=status_fn)

        if progress_fn:
            progress_fn(0.9)

        try:
            vs = Chroma(
                persist_directory=PDR_DIR,
                embedding_function=get_embeddings_model(),
                collection_name=COLLECTION_NAME_PDR,
            )
            child_count = len(vs.get()["ids"])
        except Exception:
            child_count = "?"

        if progress_fn:
            progress_fn(1.0)

        return (
            f"✅ PDR пересобран!\n"
            f"📁 Документов: {n}\n"
            f"🔍 Дочерних чанков (~850 симв.): {child_count}\n"
            f"📖 Родительские чанки (~3200 симв.) → LocalFileStore\n"
            f"📁 Новая PDR папка: {PDR_DIR}\n"
            f"📁 Новый PDR docstore: {PDR_DOCSTORE}"
        )
    finally:
        set_active_db_paths(pdr_dir=old_pdr_dir, pdr_docstore=old_docstore)
        st.cache_resource.clear()



def rebuild_pdr_in_place(status_fn=None, progress_fn=None) -> str:
    """Пересборка только PDR в рабочих папках без пересборки основной vector БД."""
    release_chroma_resources()
    for d in [PDR_DIR, PDR_DOCSTORE]:
        if os.path.exists(d):
            shutil.rmtree(d)
        os.makedirs(d, exist_ok=True)
    st.cache_resource.clear()

    files = [f for f in os.listdir(SOURCE_DOCS_DIR) if f.endswith((".txt", ".pdf", ".docx"))]
    if not files:
        return "❌ Папка source_docs пуста."

    documents = []
    total_files = len(files)
    for i, fname in enumerate(files):
        path = os.path.join(SOURCE_DOCS_DIR, fname)
        try:
            loader = get_loader(path)
            if loader:
                loaded = loader.load()
                for d in loaded:
                    d.metadata["indexed_at"] = datetime.datetime.now().isoformat()
                    d.metadata["doc_version"] = "v1"
                documents.extend(loaded)
        except Exception as e:
            if status_fn:
                status_fn(f"⚠️ {fname}: {e}")
        if progress_fn:
            progress_fn(0.05 + (i + 1) / max(total_files, 1) * 0.35)

    if not documents:
        return "❌ Не удалось загрузить файлы."

    if status_fn:
        status_fn(f"📦 Загружено {len(documents)} документов — строю только PDR...")
    if progress_fn:
        progress_fn(0.45)

    n = add_docs_to_pdr(documents, status_fn=status_fn)

    if progress_fn:
        progress_fn(0.9)

    try:
        vs = Chroma(
            persist_directory=PDR_DIR,
            embedding_function=get_embeddings_model(),
            collection_name=COLLECTION_NAME_PDR,
        )
        child_count = len(vs.get()["ids"])
    except Exception as e:
        child_count = 0
        if status_fn:
            status_fn(f"⚠️ Проверка PDR после сборки: {e}")

    if progress_fn:
        progress_fn(1.0)

    return (
        f"✅ PDR пересобран в рабочих папках!\n"
        f"📁 Документов: {n}\n"
        f"🔍 Дочерних чанков (~850 симв.): {child_count}\n"
        f"📖 Родительские чанки (~3200 симв.) хранятся в папке: {PDR_DOCSTORE}\n"
        f"📁 PDR vector: {PDR_DIR}\n"
        f"📁 PDR docstore: {PDR_DOCSTORE}"
    )

# ==========================================
# РЕРАНКИНГ С ДЕДУПЛИКАЦИЕЙ ПО ИСТОЧНИКУ
# ==========================================
def rerank_documents(query: str, context_snippet: str, docs, reranker_choice: str, top_k: int = 5):
    config = RERANKER_CONFIGS.get(reranker_choice)
    if not config:
        return [(doc, 0.0) for doc in docs[:top_k]]

    reranker = get_reranker_model(reranker_choice)
    if reranker is None:
        return [(doc, 0.0) for doc in docs[:top_k]]

    full_context = f"{query} {context_snippet}" if context_snippet else query
    full_context_trimmed = full_context[:config["max_length"] * 4]

    try:
        if config["type"] == "local_crossencoder":
            pairs = [[full_context_trimmed, doc.page_content] for doc in docs]
            base_scores = reranker.predict(pairs)
        elif config["type"] == "cohere_api":
            results = reranker.rerank(
                query=full_context_trimmed,
                documents=[d.page_content for d in docs],
                top_n=min(top_k * 2, len(docs)),
                model=config["model_name"],
            )
            base_scores = [0.0] * len(docs)
            for r in results.results:
                base_scores[r.index] = r.relevance_score
        else:
            return [(doc, 0.0) for doc in docs[:top_k]]

        # Финальный score с иерархическим весом
        scored = []
        for doc, base_score in zip(docs, base_scores):
            doc_type = detect_document_type(
                doc.metadata.get("source", ""), doc.page_content
            )
            weight = DOCUMENT_TYPE_WEIGHTS.get(doc_type, 1.0)
            final_score = base_score * weight
            doc.metadata["doc_type"] = doc_type
            doc.metadata["base_score"] = float(base_score)
            doc.metadata["hierarchy_weight"] = weight
            doc.metadata["indexed_at"] = doc.metadata.get(
                "indexed_at", datetime.datetime.now().isoformat()
            )
            scored.append((doc, final_score))

        scored.sort(key=lambda x: x[1], reverse=True)

        # Дедупликация по источнику: лучший чанк от каждого документа
        seen_sources: dict = {}
        for doc, score in scored:
            src = doc.metadata.get("source", "unknown")
            if src not in seen_sources or score > seen_sources[src][1]:
                seen_sources[src] = (doc, score)

        unique = sorted(seen_sources.values(), key=lambda x: x[1], reverse=True)
        return unique[:max(top_k, 1)]

    except Exception as e:
        st.error(f"❌ Ошибка реранкинга: {e}")
        return [(doc, 0.0) for doc in docs[:top_k]]


def get_hybrid_context(
    query: str, context_snippet: str, k: int = 7, reranker_choice: str = "BGE v2-m3 (быстрый)"
) -> tuple[str, float]:
    """
    Возвращает (context_str, max_score).
    PDR-ветка: ищет по child (~850 симв.) → отдаёт parent (~3200 симв.) в реранкер → LLM.
    Fallback: обычный BM25 + Vector (~1000 симв. чанки).
    """
    use_pdr = st.session_state.get("use_pdr", False)

    # ── PDR-ветка ────────────────────────────────────────────────────────────
    if use_pdr and pdr_has_documents():
        pdr_docs = get_pdr_docs(query, context_snippet, k=k)
        if pdr_docs:
            final_docs = rerank_documents(
                query=query,
                context_snippet=context_snippet,
                docs=pdr_docs,
                reranker_choice=reranker_choice,
                top_k=k,
            )
            max_score = max((s for _, s in final_docs), default=0.0)
            if max_score < MIN_RELEVANCE_SCORE:
                log_low_recall(query, max_score)
            context_str = ""
            for i, (doc, score) in enumerate(final_docs, 1):
                src      = os.path.basename(doc.metadata.get("source", "НПА"))
                doc_type = doc.metadata.get("doc_type", "")
                context_str += (
                    f"[Источник {i}: {src}] ({doc_type} · PDR ~3200 симв. · {score:.3f})\n"
                    f"{doc.page_content}\n---\n"
                )
            return context_str, max_score
        # PDR ничего не нашёл → fallback
        log_low_recall(query, 0.0)

    # ── Fallback: BM25 + Vector ──────────────────────────────────────────────
    vs = load_vector_db()
    if not vs:
        return (
            "⚠️ База знаний пуста. Загрузите Пленумы и Обзоры ВС РФ в разделе «Управление базой».",
            0.0,
        )

    # Jina v3: task=retrieval.query через официальный префикс
    query_with_prefix = f"retrieval.query: {query}"
    search_text = f"{query_with_prefix} {context_snippet}" if context_snippet else query_with_prefix

    bm25 = get_bm25_retriever()
    if bm25:
        vector_retriever = vs.as_retriever(search_kwargs={"k": max(k * 3, 15)})
        bm25.k = max(k * 3, 15)
        ensemble = EnsembleRetriever(
            retrievers=[bm25, vector_retriever],
            weights=[BM25_WEIGHT, VECTOR_WEIGHT],
        )
        docs = ensemble.invoke(search_text)
    else:
        docs = vs.as_retriever(search_kwargs={"k": max(k * 3, 15)}).invoke(search_text)

    if not docs:
        return "Ничего не найдено в базе Пленумов и Обзоров ВС РФ.", 0.0

    final_docs = rerank_documents(
        query=query,
        context_snippet=context_snippet,
        docs=docs,
        reranker_choice=reranker_choice,
        top_k=k,
    )
    max_score = max((s for _, s in final_docs), default=0.0)
    if max_score < MIN_RELEVANCE_SCORE:
        log_low_recall(query, max_score)

    context_str = ""
    for i, (doc, score) in enumerate(final_docs, 1):
        src      = os.path.basename(doc.metadata.get("source", "НПА"))
        doc_type = doc.metadata.get("doc_type", "")
        context_str += (
            f"[Источник {i}: {src}] ({doc_type} · ~1000 симв. · {score:.3f})\n"
            f"{doc.page_content}\n---\n"
        )
    return context_str, max_score


# ==========================================
# УПРАВЛЕНИЕ БАЗОЙ ЗНАНИЙ
# ==========================================
TEXT_SPLITTER = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP,
    separators=[
        "\nПостановление Пленума",
        "\nПункт ", "\nп. ",
        "\nСтатья ", "\nст. ", "\nГлава ",
        "\nУСТАНОВИЛ:", "\nОПРЕДЕЛИЛ:", "\nРЕШИЛ:", "\nПОСТАНОВИЛ:",
        "\n\n", "\n", " ", "",
    ],
)


def add_files_to_knowledge_base(uploaded_files):
    status_text = st.empty()
    progress_bar = st.progress(0)
    build_dir = RAG_DIR
    vectorstore = Chroma(
        persist_directory=build_dir,
        embedding_function=get_embeddings_model(),
        collection_name=COLLECTION_NAME,
    )
    documents = []
    for i, file in enumerate(uploaded_files):
        safe_name = os.path.basename(file.name)
        file_path = os.path.join(SOURCE_DOCS_DIR, safe_name)
        try:
            vectorstore.delete(where={"source": file_path})
        except Exception:
            pass
        with open(file_path, "wb") as buf:
            buf.write(file.getbuffer())
        try:
            loader = get_loader(file_path)
            if loader:
                loaded = loader.load()
                # Добавляем метаданные версионирования
                for doc in loaded:
                    doc.metadata["indexed_at"] = datetime.datetime.now().isoformat()
                    doc.metadata["doc_version"] = "v1"
                documents.extend(loaded)
        except Exception as e:
            st.error(f"Ошибка с файлом {safe_name}: {e}")
        progress_bar.progress((i + 1) / len(uploaded_files) * 0.3)
    if not documents:
        return "❌ Не удалось прочитать файлы."
    splits = TEXT_SPLITTER.split_documents(documents)
    ids = generate_doc_ids(splits)
    vectorstore.add_documents(documents=splits, ids=ids)
    st.cache_resource.clear()
    get_bm25_retriever()
    progress_bar.progress(1.0)
    return f"✅ Добавлено/обновлено фрагментов: {len(splits)}"




def release_chroma_resources():
    try:
        st.cache_resource.clear()
        st.cache_data.clear()
    except Exception:
        pass


def safe_remove_dir(path: str):
    import gc
    import time
    release_chroma_resources()
    gc.collect()
    if not os.path.exists(path):
        return
    last_error = None
    for _ in range(6):
        try:
            shutil.rmtree(path)
            return
        except PermissionError as e:
            last_error = e
            gc.collect()
            time.sleep(0.6)
        except Exception:
            raise
    raise last_error



def rotate_dir(path: str):
    import time
    if not os.path.exists(path):
        return None
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    rotated = f"{path}_old_{ts}"
    base_rotated = rotated
    n = 1
    while os.path.exists(rotated):
        rotated = f"{base_rotated}_{n}"
        n += 1
    os.rename(path, rotated)
    return rotated


def cleanup_old_db_dirs(prefixes=None):
    if prefixes is None:
        prefixes = [RAG_DIR + "_old_", PDR_DIR + "_old_", PDR_DOCSTORE + "_old_"]
    removed = []
    for name in os.listdir(BASE_DIR):
        full = os.path.join(BASE_DIR, name)
        if not os.path.isdir(full):
            continue
        if any(full.startswith(p) for p in prefixes):
            try:
                safe_remove_dir(full)
                removed.append(full)
            except Exception:
                pass
    return removed



def make_shadow_dir(base_path: str, suffix: str = "_shadow"):
    import time
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    candidate = f"{base_path}{suffix}_{ts}"
    n = 1
    while os.path.exists(candidate):
        candidate = f"{base_path}{suffix}_{ts}_{n}"
        n += 1
    os.makedirs(candidate, exist_ok=True)
    return candidate


def set_active_db_paths(rag_dir=None, pdr_dir=None, pdr_docstore=None):
    global RAG_DIR, PDR_DIR, PDR_DOCSTORE
    if rag_dir:
        RAG_DIR = rag_dir
    if pdr_dir:
        PDR_DIR = pdr_dir
    if pdr_docstore:
        PDR_DOCSTORE = pdr_docstore


def reset_active_db_paths():
    set_active_db_paths(
        os.path.join(BASE_DIR, "rag_db"),
        os.path.join(BASE_DIR, "rag_db_pdr"),
        os.path.join(BASE_DIR, "rag_db_pdr_docs"),
    )

def rebuild_knowledge_base(target_dir=None, status_fn=None, progress_fn=None):
    release_chroma_resources()
    build_dir = target_dir or make_shadow_dir(os.path.join(BASE_DIR, "rag_db"))
    os.makedirs(build_dir, exist_ok=True)

    old_rag_dir = RAG_DIR
    set_active_db_paths(rag_dir=build_dir)
    st.cache_resource.clear()

    try:
        files = [f for f in os.listdir(SOURCE_DOCS_DIR) if f.endswith((".txt", ".pdf", ".docx"))]
        if not files:
            return "❌ Папка source_docs пуста."

        documents = []
        total_files = len(files)
        for i, f in enumerate(files):
            path = os.path.join(SOURCE_DOCS_DIR, f)
            try:
                loader = get_loader(path)
                if loader:
                    loaded = loader.load()
                    for doc in loaded:
                        doc.metadata["indexed_at"] = datetime.datetime.now().isoformat()
                        doc.metadata["doc_version"] = "v1"
                    documents.extend(loaded)
            except Exception as e:
                if status_fn:
                    status_fn(f"⚠️ Не удалось загрузить {f}: {e}")
            if progress_fn:
                progress_fn(0.05 + (i + 1) / max(total_files, 1) * 0.35)

        if not documents:
            return "❌ Не удалось загрузить документы из source_docs."

        if status_fn:
            status_fn(f"✂️ Нарезаю {len(documents)} документов на фрагменты...")
        splits = TEXT_SPLITTER.split_documents(documents)
        if progress_fn:
            progress_fn(0.5)

        if status_fn:
            status_fn(f"🧠 Создаю векторную базу: {len(splits)} чанков...")
        vectorstore = Chroma(
            persist_directory=build_dir,
            embedding_function=get_embeddings_model(),
            collection_name=COLLECTION_NAME,
        )
        added = 0
        total = max(len(splits), 1)
        for i in range(0, len(splits), BATCH_SIZE):
            batch = splits[i:i + BATCH_SIZE]
            vectorstore.add_documents(documents=batch, ids=generate_doc_ids(batch))
            added += len(batch)
            if progress_fn:
                progress_fn(0.5 + (added / total) * 0.45)
            if status_fn:
                status_fn(f"💾 Векторная БД: {added}/{len(splits)} чанков...")

        st.cache_resource.clear()
        if progress_fn:
            progress_fn(1.0)
        return f"✅ База пересобрана в новой папке: {build_dir}\nВсего фрагментов: {len(splits)}"
    finally:
        set_active_db_paths(rag_dir=old_rag_dir)
        st.cache_resource.clear()


def continue_indexing(statustext=None, progressbar=None):
    if statustext is None:
        statustext = st.empty()
    if progressbar is None:
        progressbar = st.progress(0)
    try:
        vectorstore = Chroma(
            persist_directory=RAG_DIR,
            embedding_function=get_embeddings_model(),
            collection_name=COLLECTION_NAME,
        )
        existing_count = len(vectorstore.get()["ids"])
        statustext.text(f"📊 Уже проиндексировано: {existing_count} чанков")
    except Exception as e:
        st.error(f"Ошибка загрузки БД: {e}")
        return "❌ Не удалось загрузить базу"
    existing_data = vectorstore.get(include=["metadatas"])
    indexed_sources = set(
        os.path.basename(m.get("source", ""))
        for m in existing_data.get("metadatas", []) if m
    )
    all_files = [f for f in os.listdir(SOURCE_DOCS_DIR) if f.endswith((".txt", ".pdf", ".docx"))]
    new_files = [f for f in all_files if f not in indexed_sources]
    if not new_files:
        statustext.text("✅ Все файлы уже проиндексированы!")
        progressbar.progress(1.0)
        return f"✅ База актуальна! {existing_count} чанков"
    documents = []
    for i, f in enumerate(new_files):
        path = os.path.join(SOURCE_DOCS_DIR, f)
        try:
            loader = get_loader(path)
            if loader:
                loaded = loader.load()
                for doc in loaded:
                    doc.metadata["indexed_at"] = datetime.datetime.now().isoformat()
                    doc.metadata["doc_version"] = "v1"
                documents.extend(loaded)
        except Exception as e:
            st.warning(f"⚠️ {f}: {e}")
        progressbar.progress(0.2 + (i + 1) / len(new_files) * 0.2)
    if not documents:
        return "⚠️ Не удалось загрузить новые документы"
    splits = TEXT_SPLITTER.split_documents(documents)
    added = 0
    for i in range(0, len(splits), BATCH_SIZE):
        batch = splits[i:i + BATCH_SIZE]
        try:
            vectorstore.add_documents(documents=batch, ids=generate_doc_ids(batch))
            added += len(batch)
            progressbar.progress(0.4 + (added / len(splits)) * 0.6)
            statustext.text(f"💾 {added}/{len(splits)} новых чанков...")
        except Exception as e:
            st.error(f"Ошибка батча {i}: {e}")
    st.cache_resource.clear()
    final_count = len(vectorstore.get()["ids"])
    progressbar.progress(1.0)
    return (
        f"✅ Допиндексация завершена!\n"
        f"📁 Добавлено файлов: {len(new_files)}\n"
        f"📊 Было: {existing_count} → Стало: {final_count} чанков"
    )


# ==========================================
# DOCX ЭКСПОРТ
# ==========================================
def create_docx_from_text(text: str) -> BytesIO:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.0
    for line in text.split("\n"):
        if not line.strip():
            doc.add_paragraph()
        elif any(k in line.upper() for k in ["АНАЛИЗ", "РАЗДЕЛ", "ВЫВОД", "НАРУШЕНИЕ", "СООТВЕТСТВИЕ", "ТРЕБОВАНИЕ"]):
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.bold = True
        else:
            doc.add_paragraph(line)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# ==========================================
# LLM: ЕДИНЫЙ ДИСПЕТЧЕР (Gemini + OpenRouter)
# ==========================================
SYSTEM_PROMPT = """
Ты — опытный юрист-аналитик, специализирующийся на гражданских делах в судах общей юрисдикции РФ.

Твоя задача: анализировать решения суда первой инстанции и доводы апелляционных жалоб на соответствие \
позициям Пленумов Верховного Суда РФ и Обзорам судебной практики ВС РФ.

ПРИНЦИПЫ:
1. Опирайся ИСКЛЮЧИТЕЛЬНО на фрагменты из RAG-базы Пленумов и Обзоров ВС РФ, предоставленные тебе.
2. Не выдумывай реквизиты — цитируй только то, что есть в контексте.
3. При несоответствии — указывай конкретный пункт Пленума или позицию Обзора.
4. Разграничивай: (а) прямое нарушение позиции ВС РФ; (б) неприменение Пленума; (в) соответствие ВС РФ.
5. Если в базе нет релевантного Пленума — честно сообщи об этом.
6. Удерживай контекст всей беседы: уточняющие вопросы отвечай с учётом ранее обсуждённого.

СПЕЦИАЛИЗАЦИЯ: только гражданские дела (ГПК РФ, ГК РФ, ЖК РФ, СК РФ и Пленумы ВС РФ по ним).
"""


def call_llm_streaming(
    provider_name: str,
    api_key: str,
    model_id: str,
    chat_history: list,
    new_user_message: str,
    temperature: float,
    top_p: float,
    max_tokens: int,
):
    """
    Единый стриминговый вызов для Gemini и OpenRouter (OpenAI-compat).
    Yields: накопленный текст (partial response).
    """
    prov_cfg = LLM_PROVIDERS[provider_name]

    if prov_cfg["type"] == "gemini":
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(
            model_name=model_id,
            system_instruction=SYSTEM_PROMPT,
        )
        gen_config = GenerationConfig(
            temperature=temperature,
            top_p=top_p,
            max_output_tokens=max_tokens,
        )
        chat = model.start_chat(history=chat_history)
        stream = chat.send_message(
            new_user_message, generation_config=gen_config, stream=True
        )
        full = ""
        for chunk in stream:
            if chunk.text:
                full += chunk.text
                yield full

    else:  # prov_cfg["type"] == "openrouter" (OpenAI-compat)
        client = OpenAI(
            api_key=api_key,
            base_url=prov_cfg["base_url"],
            default_headers={
                "HTTP-Referer": "https://legal-analyzer.local",
                "X-Title": "Legal Analyzer v6.5",
            },
        )
        # Строим messages в формате OpenAI
        messages = [{"role": "system", "content": SYSTEM_PROMPT}]
        for msg in chat_history:
            role = "assistant" if msg["role"] == "model" else msg["role"]
            messages.append({"role": role, "content": msg["parts"][0]})
        messages.append({"role": "user", "content": new_user_message})

        stream = client.chat.completions.create(
            model=model_id,
            messages=messages,
            temperature=temperature,
            top_p=top_p,
            max_tokens=max_tokens,
            stream=True,
        )
        full = ""
        for chunk in stream:
            delta = chunk.choices[0].delta.content or ""
            full += delta
            yield full


# ==========================================
# ФОРМИРОВАНИЕ СООБЩЕНИЯ ДЛЯ LLM
# ==========================================
def build_user_message(
    user_prompt: str,
    court_decision_text: str,
    appeal_text: str,
    rag_context: str,
    rag_max_score: float,
    is_first_turn: bool,
) -> str:
    decision_block = ""
    appeal_block = ""

    if is_first_turn:
        if court_decision_text:
            decision_block = f"\n=== РЕШЕНИЕ СУДА ПЕРВОЙ ИНСТАНЦИИ ===\n{court_decision_text}\n"
        if appeal_text:
            appeal_block = f"\n=== ДОВОДЫ АПЕЛЛЯЦИОННОЙ ЖАЛОБЫ ===\n{appeal_text}\n"

    # Предупреждение о низкой релевантности
    if rag_max_score < MIN_RELEVANCE_SCORE:
        rag_warning = (
            f"\n⚠️ ВНИМАНИЕ: релевантность найденных фрагментов низкая ({rag_max_score:.3f} < {MIN_RELEVANCE_SCORE}). "
            "Возможно, по данному вопросу Пленума ВС РФ в базе нет. "
            "Укажи это явно в ответе и не додумывай.\n"
        )
    else:
        rag_warning = ""

    rag_block = f"\n=== НАЙДЕННЫЕ ПЛЕНУМЫ И ОБЗОРЫ ВС РФ ===\n{rag_warning}{rag_context}\n"
    task_block = f"\n=== ЗАДАЧА ===\n{user_prompt}\n"

    return decision_block + appeal_block + rag_block + task_block


# ==========================================
# КЭШ ОТВЕТОВ LLM
# ==========================================
import hashlib as _hashlib

LLM_RESPONSE_CACHE: dict = {}

def _make_cache_key(user_prompt: str, decision_text: str, rag_context: str, history_len: int) -> str:
    raw = f"{user_prompt}||{decision_text[:500]}||{rag_context[:800]}||{history_len}"
    return _hashlib.md5(raw.encode("utf-8")).hexdigest()

# ==========================================
# ПАРСЕР ВЫВОДОВ LLM → КОНСТРУКТОР ДОКУМЕНТА
# ==========================================

def extract_conclusions(llm_response: str) -> list:
    conclusions = []
    seen: set = set()

    def _add(text: str, ctype: str) -> None:
        import re
        text = re.sub(r'\s+', ' ', text).strip().rstrip('.')
        if len(text) < 25 or text in seen:
            return
        seen.add(text)
        conclusions.append({"text": text, "type": ctype, "selected": True})

    import re
    for m in re.finditer(r'(?:^|\n)\s*\d+[.)]\s+(.{25,400})', llm_response):
        txt = m.group(1).strip()
        if re.search(r'нарушение|не соответствует|ошибочн|не учёл|не применил|противоречит', txt, re.I):
            _add(txt, 'violation')
        elif re.search(r'соответствует|согласуется|правомерно|подтверждается', txt, re.I):
            _add(txt, 'compliance')
        else:
            _add(txt, 'general')

    for m in re.finditer(
        r'((?:нарушение|не соответствует|противоречит|суд не учёл|ошибочно)[^\.\n]{20,250}'
        r'(?:Пленум|п\.\s*\d+|ст\.\s*\d+|ГПК|ГК)[^\.\n]{0,100})',
        llm_response, re.IGNORECASE,
    ):
        _add(m.group(1), 'violation')

    for m in re.finditer(
        r'((?:соответствует|согласуется|подтверждается)[^\.\n]{20,250}'
        r'(?:Пленум|п\.\s*\d+|ст\.\s*\d+)[^\.\n]{0,100})',
        llm_response, re.IGNORECASE,
    ):
        _add(m.group(1), 'compliance')

    for m in re.finditer(
        r'((?:суд первой инстанции|судебная коллегия)[^\.\n]{15,300}\.)',
        llm_response, re.IGNORECASE,
    ):
        _add(m.group(1), 'general')

    return conclusions[:30]

# ── Вспомогательные функции для вкладки «База знаний» ──────────────────────
def split_text_to_documents(text: str, metadata: dict) -> list:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.schema import Document as _LCDoc
    splitter = RecursiveCharacterTextSplitter(chunk_size=850, chunk_overlap=120)
    return [_LCDoc(page_content=c, metadata=metadata) for c in splitter.split_text(text)]

def add_document_to_kb(text: str, metadata: dict = None) -> None:
    """Добавляет документ в индексируемую базу через стандартный пайплайн."""
    docs = split_text_to_documents(text, metadata or {})
    add_files_to_knowledge_base(docs)

def get_kb_stats() -> dict:
    """Возвращает статистику ChromaDB."""
    try:
        vdb = load_vector_db()
        return {
            "total_chunks": vdb._collection.count(),
            "embedding_model": "jina-embeddings-v3",
            "persist_directory": RAG_DIR,
        }
    except Exception as e:
        return {"error": str(e)}


# ==========================================
# ПРОМПТЫ ДЛЯ ГЕНЕРАЦИИ ДОКУМЕНТОВ
# ==========================================
SYSTEM_PROMPT_OTMENA = '<FORMATTING_RULES>\n-даты пиши в формате - 10.05.2025\n-при написании даты с месяцем буквами: пиши например "9 мая 2025 г.", а не "09 мая 2025 г"\n-числа с копейками пиши в формате - "10545,57 рублей"\n-при подготовке текста ап.определения не добавляй интервал после абзаца\n-в тексте резолютивной части ап.определения не допускается писать числа или пояснения в скобках\n- в резолютивной части ап.определения ФИО сторон нужно писать полностью, в описательной и мотивировочной части - в формате "Фамилия С.В."\n- избегай использования тире, двойного дефиса и ; (точки с запятой) в тексте\n-избегай тавтологии в анафорах абзацев (начальное слово следующих друг за другом абзацев должно отличатся)\n- дели логически завершенные мысли по абзацам. Если цитируешь текст нормы права - каждую норму приводи в новом абзаце.\n- ссылки на нормы всегда писать в формате : ч. 1 ст. 330 ГПК РФ, а не ст. 330 ч. 1 ГПК\n-слова "истец" и "ответчик" не склонять в женском роде: формы "истица, истице" и "ответчица, ответчице" недопустимы.\n</FORMATTING_RULES>\n\n# РОЛЬ И ЗАДАЧА\n\nТы опытный судья суда субъекта Российской Федерации.\n\nТвоя задача составить/написать проект апелляционного определения об ОТМЕНЕ решения суда по гражданскому делу и принятии нового решения (или изменении решения), удовлетворив апелляционную жалобу (полностью или частично) на основе решения суда первой инстанции <SOURCE_DECISION> и доводов апелляционной жалобы <APPEAL_ARGUMENTS>, придерживаясь изложенной в этом промпте структуры <OUTPUT_STRUCTURE>.\n\nНапиши мотивированное обоснование незаконности обжалуемого решения, соглашаясь с доводами апелляционной жалобы <APPEAL_ARGUMENTS> (если апелляционных жалоб больше и от разных лиц, соответственно). Обоснуй необходимость принятия нового решения со ссылками на обстоятельства дела и нормы законов, которые были нарушены судом первой инстанции.\n\nПиши в стиле судьи апелляционной инстанции.\n\n<APPEAL_ARGUMENTS>\n{appeal_arguments}\nМожешь учитывать свои собственные выводы по предыдущему вопросу в данном чате.\n</APPEAL_ARGUMENTS>\n\n<ANALYSIS_BEFORE_WRITING>\nПрежде чем писать текст определения, выполни внутренний анализ. В итоговый текст его не включай.\n\nЕСЛИ <APPEAL_ARGUMENTS> ЗАПОЛНЕН:\n\nШаг 1. Исходя из доводов <APPEAL_ARGUMENTS>, определи, какой из пунктов ч. 1 ст. 330 ГПК РФ применим (1-4). Задача - не переоценивать доводы жалобы, а найти процессуальное основание для уже обозначенного вывода.\n\nШаг 2. Выяви конкретную норму материального права, которую нарушил или не применил суд первой инстанции в той части, которая обжалуется согласно <APPEAL_ARGUMENTS>.\n\nШаг 3. Найди в <SOURCE_DECISION> факты, которые противоречат выводам суда или остались без оценки, и которые подтверждают правоту доводов из <APPEAL_ARGUMENTS>.\n\nШаг 4. Уточни объём отмены и итог нового решения строго в соответствии с <APPEAL_ARGUMENTS>: в какой части отменяется решение? что постановляется в этой части (удовлетворить / отказать)? в какой части решение остаётся без изменения (если отмена частичная)?\n\nЕСЛИ <APPEAL_ARGUMENTS> НЕ ЗАПОЛНЕН ИЛИ ПУСТОЙ:\n\nШаг 1. Самостоятельно изучи <SOURCE_DECISION> и определи, какой из пунктов ч. 1 ст. 330 ГПК РФ применим (1-4).\n\nШаг 2. Выяви конкретную норму материального или процессуального права, которую нарушил или не применил суд первой инстанции.\n\nШаг 3. Найди в <SOURCE_DECISION> факты и доказательства, которые противоречат выводам суда или остались без надлежащей правовой оценки.\n\nШаг 4. Самостоятельно определи итог нового решения: удовлетворить иск полностью / частично / отказать?\n\nВ ОБОИХ СЛУЧАЯХ: только после завершения всех шагов приступай к написанию текста определения согласно <OUTPUT_STRUCTURE>.\n</ANALYSIS_BEFORE_WRITING>\n\n# ГЛАВНЫЕ ЗАПРЕТЫ (CRITICAL CONSTRAINTS):\n\nСпецифические инструкции для каждого конкретного подраздела в структуре <OUTPUT_STRUCTURE> имеют приоритет над общими запретами и ограничениями.\n\nВАЖНО! Составляй проект апелляционного определения без вводной части и резолютивной части (где указан состав суда, какое дело рассмотрено и каков результат рассмотрения), но придерживаясь далее указанной структуры.\n\nЗаголовки "А)", "В)", "Подраздел (наименование подраздела):", теги и тому подобное в этом промпте служат только для навигации. Запрет на заголовки является абсолютным и распространяется на любые слова и фразы, выполняющие функцию заголовка или рубрики. Итоговый текст проекта апелляционного определения должен представлять собой сплошной связный текст, разбитый только на смысловые абзацы, без каких-либо заголовков, подзаголовков, рубрик и нумерации разделов.\n\nЖёсткое требование дословности: все фрагменты, воспроизводимые с пометкой "дословно", воспроизводиться БЕЗ ИЗМЕНЕНИЙ, включая пунктуацию, разбивку на абзацы и форматирование, если иное прямо не указано по тексту промпта.\n\nИспользуй классический, общепринятый стиль изложения судебных актов по гражданским делам, принятый в Российской Федерации.\n\nПиши сплошным текстом, не используя рубрикацию, буллиты, заголовки и тому подобное, сохраняй разбивку на абзацы по смыслу. Пиши достаточно пространно, но ясно, со ссылками на подходящие нормативно-правовые акты.\n\n<OUTPUT_STRUCTURE>\nСтруктура проекта апелляционного определения: проект должен содержать следующие части: А) описательная часть проекта апелляционного определения, В) мотивировочная часть проекта апелляционного определения, каждая из которых соответствует следующим условиям:\n\nА) Описательная часть:\n\nБлизко к тексту воспроизведи часть решения суда начиная со слова "установил", в том числе включая описание явки в судебное заседание сторон и лиц, участвующих в деле в суде первой инстанции, описание их объяснений и возражений.\n\nВоспроизведение части решения суда прекрати непосредственно перед переходными фразами, такими как "Обсудив вопрос о возможности рассмотрения дела...", "Выслушав объяснения явившихся лиц..." или "Исследовав материалы дела, суд приходит к следующему". Абзац с этим переходным смыслом не должен копироваться в проект.\n\nИсправляй в описательной части грамматические, орфографические, стилистические ошибки, улучшай разбивку на смысловые абзацы, преобразуя повествование в прошедшее время.\n\nВажно включай: явку/неявку в суд первой инстанции, заявленные требования и их обоснование, возражения ответчика и иных лиц, участвующих в деле, ходатайства, позиции сторон - без выводов суда, без оценки доказательств и без формул "суд установил/считает/приходит к выводу".\n\nСтрого и чётко разделяй на абзацы по смыслу. Всегда начинай с нового абзаца с абзацным отступом обоснование исковых требований, описание явки в суд первой инстанции из решения суда.\n\nЗаверши описательную часть описанием ошибочного результата по делу с прямой цитатой резолютивной части приведённого решения суда в формате "Решением.... суда (копируй наименование суда и его дату) постановлено:..." (Найди в конце решения суда <SOURCE_DECISION> текст после "Решил" и используй его без кавычек, с прописной буквы).\n\nВ) Мотивировочная часть:\n\nПодраздел (краткое изложение апелляционной жалобы):\nВАЖНО: В этом подразделе ИГНОРИРУЙ запрет на дословное воспроизведение. Обязан исправить грамматические, орфографические и стилистические ошибки в тексте жалобы, сохраняя лишь её юридический смысл.\nИзложи доводы апелляционной жалобы <APPEAL_ARGUMENTS>, начиная фразу с "В апелляционной жалобе поставлен вопрос об отмене (или изменении) решения суда как незаконного и необоснованного..."\nАкцентируй внимание на тех доводах, которые послужат основанием для отмены.\n\nПодраздел (явка лиц в суд апелляционной инстанции):\nДословно воспроизведи следующий текст, сохраняя форматирование:\n\nЛица, участвующие в деле, в судебное заседание суда апелляционной инстанции не явились о времени и месте извещались в соответствии с требованиями ст. 113 ГПК РФ, о причинах неявки суду не сообщили.\n\nИстец в судебное заседание суда апелляционной инстанции явился лично и в лице представителя по доверенности, просил апелляционную жалобу удовлетворить/оставить без удовлетворения.\n\nИстец в судебное заседание суда апелляционной инстанции не явился о времени и месте извещался в соответствии с требованиями ст. 113 ГПК РФ, о причинах неявки суду не сообщил.\n\nОтветчик в судебное заседание суда апелляционной инстанции явился лично и в лице представителя по доверенности, просил апелляционную жалобу удовлетворить/оставить без удовлетворения.\n\nОтветчик в судебное заседание суда апелляционной инстанции не явился о времени и месте извещался в соответствии с требованиями ст. 113 ГПК РФ, о причинах неявки суду не сообщил.\n\nИные лица в судебное заседание суда апелляционной инстанции не явились о времени и месте извещались в соответствии с требованиями ст. 113 ГПК РФ, о причинах неявки суду не сообщили.\n\nУточнение: ДОСЛОВНО вставьте ВЕСЬ приведенный шаблонный текст о явке. Не допускается пропуск каких-либо частей шаблона.\n\nПодраздел (процессуальные действия в апелляции):\nДословно воспроизведи следующий текст:\n\nВ соответствии с положениями ст. 167 ГПК РФ, ст. 327 ГПК РФ судебное разбирательство в суде апелляционной инстанции проведено в отсутствие не явившихся лиц, извещавшихся о времени и месте судебного заседания.\n\nСудебная коллегия, изучив материалы дела и доводы апелляционной жалобы, проверив в соответствии со статьями 327 и 327.1 Гражданского процессуального кодекса Российской Федерации законность и обоснованность решения суда первой инстанции в пределах этих доводов, исследовав имеющиеся в деле доказательства, приходит к выводу об отмене (или изменении) решения суда в силу следующего.\n\nПодраздел (правовые основания пересмотра):\nДословно воспроизведи следующий текст:\n\nВ соответствии со ст. 330 ГПК РФ основаниями для отмены или изменения решения суда в апелляционном порядке являются: 1) неправильное определение обстоятельств, имеющих значение для дела; 2) недоказанность установленных судом первой инстанции обстоятельств, имеющих значение для дела; 3) несоответствие выводов суда первой инстанции, изложенных в решении суда, обстоятельствам дела; 4) нарушение или неправильное применение норм материального права или норм процессуального права. Неправильным применением норм материального права являются: 1) неприменение закона, подлежащего применению; 2) применение закона, не подлежащего применению; 3) неправильное истолкование закона.\n\nСудебная коллегия полагает, что при рассмотрении настоящего дела судом первой инстанции такие нарушения были допущены.\n\nПодраздел (фактические обстоятельства дела и анализ ошибок суда):\n1. Укажи фактические обстоятельства, которые действительно следуют из материалов дела.\n2. Подробно опиши, в чём именно заключалась ошибка суда первой инстанции.\nИспользуй формулировки: "Суд первой инстанции не учёл, что...", "Вывод суда о том, что..., противоречит материалам дела", "Судом ошибочно применена норма...", "Вместе с тем судебная коллегия не может согласиться с данным выводом суда, поскольку...".\nОбязательно используй доводы из <APPEAL_ARGUMENTS> для обоснования порочности решения.\n\nПодраздел (оценка доводов жалобы):\nИспользуй формулировки: "Доводы апелляционной жалобы о том, что..., заслуживают внимания", "Судебная коллегия соглашается с аргументами жалобы относительно...", "Указание заявителя жалобы на нарушение судом норм процессуального права является обоснованным".\n\nПодраздел (новое правовое обоснование):\nСамостоятельно определи нормы права, подлежащие применению к существу спора. Приведи применимые статьи ГК РФ, ГПК РФ, ЖК РФ и иных федеральных законов, а также разъяснения Пленума Верховного Суда РФ.\nПрименительно к доводам о взыскании судебных расходов обязательно используй нормы ст. 88, 94, 98 ГПК РФ, а также разъяснения постановления Пленума ВС РФ от 21.01.2016 N 1.\nОбъясни, как выявленные нормы должны применяться к установленным фактам.\n\nПодраздел (выводы судебной коллегии по существу спора):\nСформулируй итоговый вывод по делу, который ложится в основу нового решения.\nПример: "При таких обстоятельствах судебная коллегия приходит к выводу о наличии правовых оснований для удовлетворения заявленных требований..."\n\nПодраздел (итоговая мотивировка отмены):\nЗаверши мотивировочную часть (адаптируй под контекст):\nУчитывая изложенное, решение суда первой инстанции нельзя признать законным и обоснованным, оно подлежит отмене на основании ст. 330 ГПК РФ с принятием по делу нового решения об удовлетворении (или отказе в удовлетворении) заявленных требований.\nЕсли отмена частичная - укажи явно, в какой части решение отменяется, а в какой остаётся без изменения.\n</OUTPUT_STRUCTURE>\n\n<FINAL_REVIEW_CHECKLIST>\nПеред выдачей результата убедись, что логика ведёт именно к ОТМЕНЕ решения. Проверь отсутствие фраз "суд правильно установил", "оснований для отмены не имеется".\nУбедись, что в проекте чётко прописано, какие именно ошибки допустил суд первой инстанции.\nУбедись, что приведено новое решение спора по существу с правовым обоснованием.\nВывери орфографию и юридическую терминологию.\nОбнаруженные слова "полагает" (в контексте автора), "мной", "мне", "меня", "я", "мои" и тому подобные перефразируй в безличные конструкции.\nВажно: в готовом проекте определения не должно быть заголовков "А)", "Подраздел" - только сплошной текст.\n</FINAL_REVIEW_CHECKLIST>\n\n{rag_context_block}\n\n<SOURCE_DECISION>\n{source_decision}\n</SOURCE_DECISION>'
SYSTEM_PROMPT_BI = '<FORMATTING_RULES>\n-даты пиши в формате - 10.05.2025\n-при написании даты с месяцем буквами: пиши например "9 мая 2025 г.", а не "09 мая 2025 г"\n-числа с копейками пиши в формате - "10545,57 рублей"\n-при подготовке текста ап.определения не добавляй интервал после абзаца\n-в тексте резолютивной части ап.определения не допускается писать числа или пояснения в скобках\n- в резолютивной части ап.определения ФИО сторон нужно писать полностью, в описательной и мотивировочной части - в формате "Фамилия С.В."\n- избегай использования тире и ; (точки с запятой) в тексте\n-избегай тавтологии в анафорах абзацев (начальное слово следующих друг за другом абзацев должно отличатся)\n- дели логически завершенные мысли по абзацам. Если цитируешь текст нормы права - каждую норму приводи в новом абзаце.\n</FORMATTING_RULES>\n\n# РОЛЬ И ЗАДАЧА\n\nТы опытный судья суда субъекта Российской Федерации.\n\nТвоя задача составить/написать проект апелляционного определения об оставлении решения суда по гражданскому делу без изменения, а апелляционную жалобу без удовлетворения на основе решения суда первой инстанции <SOURCE_DECISION> и доводов апелляционной жалобы <APPEAL_ARGUMENTS>, придерживаясь изложенной в этом промпте структуры <OUTPUT_STRUCTURE>.\n\nНапиши мотивированное, со ссылками на нормативно-правовые акты, разъяснения высших судов, отклонение доводов апелляционной жалобы <APPEAL_ARGUMENTS> со ссылками на обстоятельства дела, применимые к ситуации законы, разъяснения Верховного суда Российской Федерации.\n\nИспользуй идеи для отклонения апелляционной жалобы (если указаны, если нет - самостоятельно мотивируй), указанные в <JUDGE_LOGIC>, пиши в стиле судьи апелляционной инстанции.\n\n<APPEAL_ARGUMENTS>\n{appeal_arguments}\nМожешь учитывать свои собственные выводы по доводам апелляционной жалобы в данном чате.\n</APPEAL_ARGUMENTS>\n\n<JUDGE_LOGIC>\n{judge_logic}\n</JUDGE_LOGIC>\n\n# ГЛАВНЫЕ ЗАПРЕТЫ (CRITICAL CONSTRAINTS):\n\nСпецифические инструкции для каждого конкретного подраздела в структуре <OUTPUT_STRUCTURE> имеют приоритет над общими запретами и ограничениями.\n\nВАЖНО! Составляй проект апелляционного определения без вводной части и резолютивной части (где указан состав суда, какое дело рассмотрено и какое результат рассмотрения), но придерживаясь далее указанной структуре.\n\nСтрожайше запрещено включать в итоговый ответ названия разделов из структуры промта. Эти названия служат исключительно для понимания логики. Итоговый текст - только связный текст, разделенный на абзацы.\n\nЖёсткое требование дословности: все цитаты, выделенные в промте кавычками или с пометкой "дословно" (кроме Подраздела краткого изложения апелляционной жалобы), должны воспроизводиться БЕЗ ИЗМЕНЕНИЙ, включая пунктуацию, разбивку на абзацы, форматирование.\n\nИспользуй классический, общепринятый стиль изложения судебных актов по гражданским делам, принятый в Российской Федерации.\n\nПиши сплошным текстом, не используя рубрикацию, буллиты, заголовки и тому подобное, сохраняй разбивку на абзацы по смыслу. Пиши достаточно пространно, но ясно.\n\n<OUTPUT_STRUCTURE>\nСтруктура проекта апелляционного определения: проект должен содержать следующие части: А) описательная часть проекта апелляционного определения, В) мотивировочная часть проекта апелляционного определения.\n\nА) Описательная часть:\n\nБлизко к тексту, воспроизведи часть решения суда, после со слова "установил" в том числе включая описание явки в судебное заседание сторон и лиц, участвующих в деле в суда первой инстанции, описание их объяснений, возражений.\n\nВоспроизведение части решения суда прекрати непосредственно перед переходными фразами, такими как "Обсудив вопрос о возможности рассмотрения дела...", "Выслушав объяснения явившихся лиц..." или "Исследовав материалы дела, суд приходит к следующему".\n\nИсправляй в описательной части грамматические, орфографические, стилистические ошибки, улучшай разбивку на смысловые абзацы, преобразуя повествование в прошедшее время.\n\nВажно включай: явку/неявку в суд первой инстанции, заявленные требования и их обоснование, возражения ответчика и иных лиц участвующих в деле, ходатайства, позиции сторон - без выводов суда, без оценки доказательств и без формул "суд установил/считает/приходит к выводу".\n\nСтрого, четко разделяй на абзацы по смыслу.\n\nЗаверши описательную часть А) описанием результата по делу с прямой цитатой резолютивной части приведенного решения суда в формате "Решением.... суда (копируй наименование суда и его дату) постановлено:..." (Найди в конце решения суда <SOURCE_DECISION> текст после "Решил" и используй его без кавычек и с прописной буквы).\n\nВ) Мотивировочная часть:\n\nПодраздел (краткое изложение апелляционной жалобы):\nВАЖНО: В этом подразделе ИГНОРИРУЙ запрет на дословное цитирование. Обязан исправить грамматические, орфографические и стилистические ошибки в тексте жалобы, сохраняя лишь её юридический смысл.\nИзложи доводы апелляционной жалобы <APPEAL_ARGUMENTS>, начиная фразу с "В апелляционной жалобе, поставлен вопрос об отмене решения суда, как незаконного и необоснованного...."\n\nПодраздел (явка лиц в суд апелляционной инстанции):\nДословно воспроизведи следующий текст, сохраняя форматирование:\n\nЛица, участвующие в деле, в судебное заседание суда апелляционной инстанции не явились о времени и месте извещались в соответствии с требованиями ст. 113 ГПК РФ, о причинах неявки суду не сообщили.\n\nИстец в судебное заседание суда апелляционной инстанции явился в лице представителя по доверенности, просил апелляционную жалобу удовлетворить/оставить без удовлетворения.\n\nИстец в судебное заседание суда апелляционной инстанции не явился о времени и месте извещался в соответствии с требованиями ст. 113 ГПК РФ, о причинах неявки суду не сообщил.\n\nОтветчик в судебное заседание суда апелляционной инстанции не явился о времени и месте извещался в соответствии с требованиями ст. 113 ГПК РФ, о причинах неявки суду не сообщил.\n\nОтветчик в судебное заседание суда апелляционной инстанции явился лично и в лице представителя по доверенности, просил апелляционную жалобу удовлетворить/оставить без удовлетворения.\n\nИные лица в судебное заседание суда апелляционной инстанции не явились о времени и месте извещались в соответствии с требованиями ст. 113 ГПК РФ, о причинах неявки суду не сообщили.\n\nУточнение: ДОСЛОВНО вставьте ВЕСЬ приведенный шаблонный текст о явке, включив все возможные варианты.\n\nПодраздел (процессуальные действия в апелляции):\nДословно воспроизведи следующий текст:\n\nВ соответствии с положениями ст. 167 ГПК РФ, ст. 327 ГПК РФ судебное разбирательство в суде апелляционной инстанции проведено в отсутствие не явившихся лиц, извещавшихся о времени и месте судебного заседания.\n\nСудебная коллегия, проверив в соответствии со статьями 327 и 327.1 Гражданского процессуального кодекса Российской Федерации законность и обоснованность решения суда первой инстанции, изучив материалы дела, доводы апелляционной жалобы, исследовав имеющиеся в деле доказательства, пришла к следующим выводам.\n\nПодраздел (правовые основания пересмотра):\nДословно воспроизведи следующий текст:\n\nВ соответствии со ст. 330 ГПК РФ основаниями для отмены или изменения решения суда в апелляционном порядке являются: 1) неправильное определение обстоятельств, имеющих значение для дела; 2) недоказанность установленных судом первой инстанции обстоятельств, имеющих значение для дела; 3) несоответствие выводов суда первой инстанции, изложенных в решении суда, обстоятельствам дела; 4) нарушение или неправильное применение норм материального права или норм процессуального права. Неправильным применением норм материального права являются: 1) неприменение закона, подлежащего применению; 2) применение закона, не подлежащего применению; 3) неправильное истолкование закона.\n\nУказанных обстоятельств, которые могли бы послужить основанием к отмене либо изменению обжалуемого решения суда, при рассмотрении дела не установлено.\n\nСуд первой инстанции при рассмотрении гражданского дела правильно определил юридически значимые обстоятельства, применил подлежащие применению нормы материального права, и постановил законное и обоснованное решение, не допустив существенных нарушений норм процессуального права.\n\nПодраздел (фактические обстоятельства дела):\nНайди в тексте приведенного решения, часть решения суда о том, что установлено судом (начинается со слов "как следует из материалов дела" или "Судом установлено" и тому подобное) и вставь в проект апелляционного определения.\nУдели особое внимание и отрази информацию о наличии обременений (залог, арест), прав третьих лиц на спорное имущество, если таковые упоминаются в решении суда.\nОценку доказательств описывай от третьего лица.\n\nПодраздел (анализ применения права):\nПриведи необходимые цитаты законов и постановлений Пленумов Верховного суда Российской Федерации, обзоров судебной практики ВС РФ и других НПА из решения суда.\n\nПодраздел (выводы суда первой инстанции):\nОбязательно должен быть "главный абзац" - выводы суда первой инстанции примерно в следующем виде: "Разрешая спор по существу, суд первой инстанции, установив, что ....... руководствуясь приведенным нормами материального права, исходил из того, что...."\nИспользуй также формулировки: "Отклоняя доводы... суд первой инстанции правомерно указывал...", "Суд первой инстанции также правильно учел, что".\nУбедись, что по всем заявленным требованиям есть ответ в главном абзаце.\n\nПо делам о защите прав потребителя при наличии соответствующих требований используй формулировки:\n"Разрешая спор в части требований о компенсации морального вреда, суд первой инстанции в соответствии с положениями ст. 15 Закона РФ от 7 февраля 1992 г. N 2300-1 "О защите прав потребителей", полагал его подлежащим компенсации с учетом нарушения ответчиком прав потребителя, при этом размер компенсации определен с учетом характера, степени причиненных нравственных страданий, требований разумности и справедливости."\n"Взысканный судом на основании п. 6 ст. 13 Закона Российской Федерации "О защите прав потребителей" штраф является мерой ответственности за несоблюдение в добровольном порядке удовлетворения требований потребителя, оснований для изменения размера потребительского штрафа, судебная коллегия не находит, полагая, что размер штрафа является соразмерным последствиям нарушенного ответчиком обязательства."\n\nПодраздел (анализ выводов суда первой инстанции):\nПриведи цитату: "С указанными выводами судебная коллегия соглашается, поскольку считает их основанным на правильном применении норм материального и процессуального права, регулирующих спорные правоотношения, с учетом установленных обстоятельств."\nПроанализируй выводы суда и доводы апелляционной жалобы, дополнительно укажи, почему выводы суда являются правильными.\nНапиши абзац усиливающий убедительность всего проекта апелляционного определения со ссылками на обстоятельства дела, нормы законов.\n\nПодраздел (Оценка доводов жалобы):\nВсегда начинай подраздел со слов "Доводы апелляционной жалобы..."\nЕсли текст апелляционной жалобы предоставлен: последовательно и мотивированно отклони каждый довод, используя формулировки: "Довод жалобы о том, что ..., является несостоятельным, поскольку...", "Ссылка в жалобе на ... не может быть принята во внимание, так как...".\nЕсли текст жалобы НЕ предоставлен, используй: "...по существу сводятся к несогласию с выводами суда первой инстанции, переоценке установленных судом обстоятельств и собранных по делу доказательств, которым суд дал надлежащую правовую оценку в решении в соответствии со статьей 67 ГПК РФ. Эти доводы не содержат фактов, которые не были бы проверены и не учтены судом первой инстанции при рассмотрении дела и имели бы юридическое значение для вынесения судебного акта по существу, влияли на обоснованность и законность судебного решения, либо опровергали выводы суда первой инстанции, в связи с чем, признаются судебной коллегией несостоятельными и не могут служить основанием для отмены решения суда."\n\nЗаверши подраздел прямой цитатой:\n"Иные доводы апелляционной жалобы сводятся к несогласию с постановленным по делу решением суда и не содержат оснований, предусмотренных ст. 330 ГПК РФ, для отмены или изменения решения суда.\n\nНарушений норм процессуального права, являющихся в силу ч. 4 ст. 330 ГПК РФ основанием для отмены принятого судебного акта, судом апелляционной инстанции не установлено.\n\nТаким образом, не имеется оснований, установленных ст. 330 ГПК РФ, для отмены или изменения решения суда по заявленным доводам."\n</OUTPUT_STRUCTURE>\n\n<FINAL_REVIEW_CHECKLIST>\nМеханизм проверки: перед отправкой ответа сравни каждую цитату из промпта с текстом в проекте определения. Убедись, что совпадение составляет 100%.\nПроанализируй текст решения на предмет фактов и суждений суда первой инстанции, которые не попали в проект, но обосновывают его выводы.\nНе ссылайся на источники, которые не можешь точно идентифицировать.\nПроверь логику своего вывода, убедись, что все подразделы имеются в проекте.\nВывери орфографию, пунктуацию, логичность изложения, юридические ошибки, противоречия.\nОбнаруженные слова "полагает", "мной", "мне", "меня", "я", "мои" и тому подобные перефразируй исходя из контекста.\nВажно: в готовом проекте определения не должно быть заголовков "А)", "Подраздел" - только сплошной текст с абзацами по смыслу.\n</FINAL_REVIEW_CHECKLIST>\n\n{rag_context_block}\n\n<SOURCE_DECISION>\n{source_decision}\n</SOURCE_DECISION>'

# ==========================================
# ИНТЕРФЕЙС
# ==========================================

# ── Боковая панель ─────────────────────────────────────────────────────────────

# ==========================================
# ЭКСПОРТ / ИМПОРТ БАЗЫ (ADMIN)
# ==========================================
def export_knowledge_base_zip(zip_path: str = "backup.zip") -> str:
    folders = [RAG_DIR, PDR_DIR, PDR_DOCSTORE, SOURCE_DOCS_DIR]
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for folder in folders:
            if os.path.exists(folder):
                for root, _, files in os.walk(folder):
                    for file in files:
                        full = os.path.join(root, file)
                        arc = os.path.relpath(full, BASE_DIR)
                        zf.write(full, arc)
    return zip_path


def import_knowledge_base_zip(uploaded_file) -> str:
    backup_path = os.path.join(BASE_DIR, "_backup_import_tmp.zip")
    with open(backup_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    release_chroma_resources()
    for folder in [RAG_DIR, PDR_DIR, PDR_DOCSTORE, SOURCE_DOCS_DIR]:
        if os.path.exists(folder):
            rotate_dir(folder)
        os.makedirs(folder, exist_ok=True)
    with zipfile.ZipFile(backup_path, "r") as zf:
        zf.extractall(BASE_DIR)
    try:
        os.remove(backup_path)
    except Exception:
        pass
    st.cache_resource.clear()
    st.cache_data.clear()
    return "✅ База успешно импортирована из backup.zip"


def render_admin_tab():
    st.subheader("🛠️ Администрирование базы знаний")

    col_upload, col_ops = st.columns(2)

    with col_upload:
        st.markdown("**1. Импорт документов**")
        uploaded_laws = st.file_uploader(
            "Загрузить PDF / TXT / DOCX",
            type=["txt", "pdf", "docx"],
            accept_multiple_files=True,
            key="admin_docs_uploader",
        )
        n_files = len(os.listdir(SOURCE_DOCS_DIR)) if os.path.exists(SOURCE_DOCS_DIR) else 0
        st.info(f"Файлов в source_docs: {n_files}")
        if st.button("📥 Добавить в базу", type="primary", use_container_width=True, key="admin_add_btn"):
            if uploaded_laws:
                with st.spinner("Индексация..."):
                    res = add_files_to_knowledge_base(uploaded_laws)
                st.success(res)
            else:
                st.warning("Сначала выберите файлы")

    with col_ops:
        st.markdown("**2. Операции с базой**")
        st.caption("rag_db_pdr хранит child-векторы для поиска, а rag_db_pdr_docs хранит parent-документы, которые возвращает PDR как большие фрагменты.")
        if st.button("♻️ Допиндексация (новые файлы)", use_container_width=True, key="admin_continue_btn"):
            status = st.empty()
            prog = st.progress(0)
            with st.spinner("Проверяю новые файлы..."):
                res = continue_indexing(status, prog)
            st.success(res)

        if st.button("🧱 Пересобрать базу полностью", use_container_width=True, key="admin_rebuild_btn"):
            try:
                with st.spinner("Освобождаю ресурсы Chroma..."):
                    release_chroma_resources()
                new_rag_dir = make_shadow_dir(os.path.join(BASE_DIR, "rag_db"))
                new_pdr_dir = make_shadow_dir(os.path.join(BASE_DIR, "rag_db_pdr"))
                new_pdr_docs = make_shadow_dir(os.path.join(BASE_DIR, "rag_db_pdr_docs"))
                with st.spinner("Полная пересборка базы в новой папке..."):
                    res_v = rebuild_knowledge_base(target_dir=new_rag_dir)
                with st.spinner("Пересборка PDR в новой папке..."):
                    res_p = rebuild_pdr(target_pdr_dir=new_pdr_dir, target_docstore=new_pdr_docs)
                st.success(res_v)
                st.success(res_p)
                st.warning(
                    "Новая база собрана в shadow-папках без удаления старой активной базы. "
                    "Чтобы начать использовать новую базу, закройте приложение, переименуйте старые папки rag_db / rag_db_pdr / rag_db_pdr_docs вручную и замените их новыми shadow-папками."
                )
                st.code(
                    f"Новая vector БД: {new_rag_dir}\n"
                    f"Новая PDR БД: {new_pdr_dir}\n"
                    f"Новый PDR docstore: {new_pdr_docs}"
                )
            except PermissionError as e:
                st.error(
                    f"❌ Даже shadow-сборка не смогла стартовать. Текст ошибки: {e}"
                )

        if st.button("🧩 Пересобрать только PDR", use_container_width=True, key="admin_rebuild_pdr_only_btn"):
            st.info("Основная vector БД rag_db не трогается. Будут пересобраны только rag_db_pdr и rag_db_pdr_docs.")
            status_pdr_only = st.empty()
            prog_pdr_only = st.progress(0)
            with st.spinner("Пересобираю только PDR..."):
                res_pdr_only = rebuild_pdr_in_place(
                    status_fn=status_pdr_only.text,
                    progress_fn=prog_pdr_only.progress,
                )
            st.success(res_pdr_only)

    st.divider()
    st.markdown("**3. Резервная копия**")
    col_exp, col_imp = st.columns(2)

    with col_exp:
        if st.button("📦 Создать backup.zip", use_container_width=True, key="admin_export_btn"):
            with st.spinner("Архивирую базу..."):
                zip_path = export_knowledge_base_zip()
            st.session_state["_backup_ready"] = zip_path
            st.success(f"✅ Архив готов ({os.path.getsize(zip_path) // 1024} КБ)")
        if st.session_state.get("_backup_ready") and os.path.exists(st.session_state["_backup_ready"]):
            with open(st.session_state["_backup_ready"], "rb") as f:
                st.download_button(
                    "⬇️ Скачать backup.zip",
                    data=f,
                    file_name="backup.zip",
                    mime="application/zip",
                    use_container_width=True,
                    key="admin_download_btn",
                )

    with col_imp:
        uploaded_backup = st.file_uploader(
            "Импортировать backup.zip",
            type=["zip"],
            key="admin_backup_uploader",
        )
        if st.button("📤 Восстановить из backup.zip", use_container_width=True, key="admin_import_btn"):
            if uploaded_backup is None:
                st.warning("Выберите файл backup.zip")
            else:
                with st.spinner("Восстановление базы..."):
                    msg = import_knowledge_base_zip(uploaded_backup)
                st.success(msg)
                st.rerun()

    st.divider()
    st.markdown("**Состояние базы:**")
    try:
        vdb = load_vector_db()
        total_chunks = len(vdb.get()["ids"]) if vdb else 0
    except Exception:
        total_chunks = 0
    st.json({
        "source_docs_files": len(os.listdir(SOURCE_DOCS_DIR)) if os.path.exists(SOURCE_DOCS_DIR) else 0,
        "rag_chunks": total_chunks,
        "pdr_ready": pdr_has_documents(),
        "rag_dir": RAG_DIR,
        "pdr_dir": PDR_DIR,
    })


with st.sidebar:
    render_api_keys_section()
    st.divider()

    # ── Реранкер ─────────────────────────────────────────────────────────
    st.header("🔎 Реранкер")
    _avail_rerankers = get_available_rerankers()
    if "reranker_choice" not in st.session_state or st.session_state["reranker_choice"] not in _avail_rerankers:
        st.session_state["reranker_choice"] = _avail_rerankers[0]
    st.session_state["reranker_choice"] = st.selectbox(
        "Модель реранкинга",
        _avail_rerankers,
        index=_avail_rerankers.index(st.session_state["reranker_choice"]),
        key="sidebar_reranker_select",
    )
    _rcfg = {}
    for _k, _v in RERANKER_CONFIGS.items():
        if _k == st.session_state["reranker_choice"]:
            _rcfg = _v
            break
    if _rcfg.get("description"):
        st.caption(_rcfg["description"])
    st.divider()

    st.header("⚙️ Настройки")
    model_choice = st.selectbox(
        "Модель",
        ["gemini-2.5-flash-preview-05-20", "gemini-2.0-flash", "openrouter"],
        key="model_choice",
    )
    if st.session_state.get("model_choice") == "openrouter":
        st.text_input(
            "OpenRouter модель",
            value=st.session_state.get("or_model", "google/gemini-2.5-flash"),
            key="or_model",
        )
    st.slider("Температура (чат)", 0.0, 1.0, 0.15, 0.05, key="temp_chat")
    st.slider("Токены ответа", 1000, 8000, 4000, 500, key="max_ctx")
    st.slider("Топ-K документов", 3, 15, 7, 1, key="top_k")
    st.divider()

    st.subheader("📄 Решение суда")
    uploaded_decision = st.file_uploader(
        "Загрузить решение (.txt/.docx/.pdf)",
        type=["txt", "docx", "pdf"],
        key="decision_file",
    )
    if uploaded_decision:
        decision_text = read_file_content(uploaded_decision)
        st.session_state["decision_text"] = decision_text
        chars = len(decision_text)
        st.markdown(f":{'green' if chars > 1000 else 'orange'}[✓ {chars:,} симв.]")
    elif "decision_text" not in st.session_state:
        st.session_state["decision_text"] = ""

    st.subheader("📋 Апелляционная жалоба")
    uploaded_appeal = st.file_uploader(
        "Загрузить жалобу (.txt/.docx/.pdf)",
        type=["txt", "docx", "pdf"],
        key="appeal_file",
    )
    if uploaded_appeal:
        appeal_text_loaded = read_file_content(uploaded_appeal)
        st.session_state["appeal_text"] = appeal_text_loaded
        chars = len(appeal_text_loaded)
        st.markdown(f":{'green' if chars > 200 else 'orange'}[✓ {chars:,} симв.]")
    elif "appeal_text" not in st.session_state:
        st.session_state["appeal_text"] = ""

    st.divider()
    if st.button("🗑️ Очистить чат", use_container_width=True):
        st.session_state["chat_history"] = []
        st.session_state["gemini_history"] = []
        st.session_state.pop("last_analysis", None)
        st.session_state.pop("draft_request", None)
        st.rerun()

# ── Три вкладки ────────────────────────────────────────────────────────────────
tab_chat, tab_docs, tab_db, tab_admin = st.tabs([
    "⚖️ Анализ дела",
    "📝 Подготовка документов",
    "📚 База знаний",
    "🛠️ Админ",
])

# ══════════════════════════════════════════════════════════════════════════
# ВКЛАДКА 1: ЧАТ / АНАЛИЗ
# ══════════════════════════════════════════════════════════════════════════
with tab_chat:
    has_decision = bool(st.session_state.get("decision_text", "").strip())
    if not has_decision:
        st.info("ℹ️ Загрузите решение суда в боковой панели для полноценного анализа.")

    col1, col2, col3 = st.columns(3)
    quick_prompt = None
    with col1:
        if st.button("🔍 Полный анализ", disabled=not has_decision, use_container_width=True):
            quick_prompt = (
                "Проведи полный правовой анализ решения суда: выяви нарушения норм материального "
                "и процессуального права, несоответствие разъяснениям Пленума ВС РФ, "
                "процессуальные ошибки и перспективы апелляционного обжалования."
            )
    with col2:
        if st.button("📋 Оценить жалобу", disabled=not has_decision, use_container_width=True):
            quick_prompt = (
                "Оцени доводы апелляционной жалобы: насколько они обоснованы, "
                "соответствуют ли разъяснениям Пленума ВС РФ, какие из них наиболее перспективны."
            )
    with col3:
        if st.button("📌 Тезисы для жалобы", disabled=not has_decision, use_container_width=True):
            quick_prompt = (
                "Извлеки ключевые нарушения из решения суда и сформулируй тезисы для "
                "апелляционной жалобы со ссылками на нормы ГПК РФ и Пленумы ВС РФ."
            )

    if "chat_history" not in st.session_state:
        st.session_state["chat_history"] = []
    if "gemini_history" not in st.session_state:
        st.session_state["gemini_history"] = []

    chat_container = st.container(height=520)
    with chat_container:
        for msg in st.session_state["chat_history"]:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

    # ── Конструктор документа ────────────────────────────────────────────────
    if st.session_state.get("last_analysis"):
        with st.expander("📋 Конструктор документа — выбери выводы для проекта", expanded=False):
            conclusions = extract_conclusions(st.session_state["last_analysis"])
            if conclusions:
                selected_texts = []
                for i, c in enumerate(conclusions):
                    icon = "🔴" if c["type"] == "violation" else ("🟢" if c["type"] == "compliance" else "📌")
                    label = f"{icon} {c['text'][:130]}{'...' if len(c['text']) > 130 else ''}"
                    if st.checkbox(label, value=True, key=f"concl_{i}"):
                        selected_texts.append(c["text"])

                user_additions = st.text_area(
                    "✏️ Ваши дополнения / уточнения к документу:",
                    height=90,
                    key="doc_additions",
                )
                doc_type_choice = st.radio(
                    "Тип документа:",
                    ["Апелляционное определение об ОТМЕНЕ",
                     "Апелляционное определение об ОСТАВЛЕНИИ БЕЗ ИЗМЕНЕНИЯ"],
                    horizontal=True,
                    key="doc_type_choice",
                )
                if st.button("▶ Передать в подготовку документа", type="primary", key="to_doc_tab"):
                    st.session_state["draft_request"] = {
                        "selected_conclusions": selected_texts,
                        "additions": user_additions,
                        "doc_type": doc_type_choice,
                    }
                    st.success("✅ Выводы переданы. Перейдите на вкладку «Подготовка документов».")
            else:
                st.info("Структурированные выводы не найдены. Задайте вопрос с перечнем нарушений.")

    # ── Ввод ─────────────────────────────────────────────────────────────────
    user_input = st.chat_input("Введите вопрос или задачу...", key="chat_input")
    if quick_prompt:
        user_input = quick_prompt

    if user_input:
        st.session_state["chat_history"].append({"role": "user", "content": user_input})
        with chat_container:
            with st.chat_message("user"):
                st.markdown(user_input)

        decision_text = st.session_state.get("decision_text", "")
        appeal_text   = st.session_state.get("appeal_text", "")
        rag_context   = ""
        try:
            rag_context = get_hybrid_context(user_input, top_k=st.session_state.get("top_k", 7))
        except Exception as e:
            st.warning(f"RAG недоступен: {e}")

        full_user_msg = build_user_message(user_input, decision_text, appeal_text, rag_context)

        cache_key = _make_cache_key(
            user_input, decision_text, rag_context,
            len(st.session_state["gemini_history"]),
        )
        if cache_key in LLM_RESPONSE_CACHE:
            answer = LLM_RESPONSE_CACHE[cache_key]
            with chat_container:
                with st.chat_message("assistant"):
                    st.markdown(answer)
                    st.caption("⚡ Из кэша")
        else:
            with chat_container:
                with st.chat_message("assistant"):
                    answer = call_llm_streaming(
                        full_user_msg,
                        st.session_state["gemini_history"],
                        temperature=st.session_state.get("temp_chat", 0.15),
                        max_tokens=st.session_state.get("max_ctx", 4000),
                    )
            LLM_RESPONSE_CACHE[cache_key] = answer

        st.session_state["chat_history"].append({"role": "assistant", "content": answer})
        st.session_state["last_analysis"] = answer
        st.session_state["gemini_history"].append({
            "role": "user", "parts": [{"text": full_user_msg}],
        })
        st.session_state["gemini_history"].append({
            "role": "model", "parts": [{"text": answer}],
        })
        st.rerun()

# ══════════════════════════════════════════════════════════════════════════
# ВКЛАДКА 2: ПОДГОТОВКА ДОКУМЕНТОВ
# ══════════════════════════════════════════════════════════════════════════
with tab_docs:
    st.subheader("📝 Генерация проекта апелляционного определения")

    decision_text_full = st.session_state.get("decision_text", "")
    appeal_text_full   = st.session_state.get("appeal_text", "")

    dcol, acol = st.columns(2)
    with dcol:
        d_chars = len(decision_text_full)
        d_color = "green" if d_chars > 1000 else ("orange" if d_chars > 0 else "red")
        st.markdown(f":{d_color}[📄 Решение: {'загружено (' + str(d_chars) + ' симв.)' if d_chars else 'не загружено'}]")
    with acol:
        a_chars = len(appeal_text_full)
        a_color = "green" if a_chars > 200 else ("orange" if a_chars > 0 else "red")
        st.markdown(f":{a_color}[📋 Жалоба: {'загружена (' + str(a_chars) + ' симв.)' if a_chars else 'не загружена'}]")

    st.divider()

    draft = st.session_state.get("draft_request", {})
    default_idx = 0
    if draft.get("doc_type", "").startswith("Апелляционное определение об ОСТАВЛЕНИИ"):
        default_idx = 1

    doc_type_tab = st.radio(
        "Тип апелляционного определения:",
        ["ОТМЕНА решения", "ОСТАВЛЕНИЕ без изменения (БИ)"],
        horizontal=True,
        key="doc_type_tab",
        index=default_idx,
    )
    is_otmena = doc_type_tab.startswith("ОТМЕНА")

    prefilled_conclusions = ""
    if draft.get("selected_conclusions"):
        st.success(f"✅ Принято {len(draft['selected_conclusions'])} вывод(а/ов) из Конструктора")
        with st.expander("Просмотр принятых выводов"):
            for txt in draft["selected_conclusions"]:
                st.markdown(f"- {txt}")
        prefilled_conclusions = "\n".join(draft["selected_conclusions"])

    extra_instructions = st.text_area(
        "✏️ Дополнительные указания / уточнения:",
        value=draft.get("additions", ""),
        height=100,
        key="extra_instructions_doc",
        help="Укажите конкретные нарушения, суммы, нормы права или иные детали.",
    )

    judge_logic_text = ""
    if not is_otmena:
        judge_logic_text = st.text_area(
            "💡 Идеи для отклонения доводов жалобы (только для БИ):",
            height=100,
            key="judge_logic_text",
            help="Укажите, почему доводы жалобы несостоятельны. Если пусто — ИИ обоснует самостоятельно.",
        )

    temp_doc = st.slider("Температура (документ)", 0.0, 0.5, 0.10, 0.02, key="temp_doc")
    st.divider()

    generate_disabled = not bool(decision_text_full.strip())
    if generate_disabled:
        st.warning("⚠️ Загрузите решение суда в боковой панели.")

    if st.button("▶ Сгенерировать проект определения", type="primary",
                 disabled=generate_disabled, key="gen_doc_btn"):

        appeal_block = ""
        if prefilled_conclusions:
            appeal_block += f"\n\nВыводы из анализа (выбраны в Конструкторе):\n{prefilled_conclusions}"
        if appeal_text_full.strip():
            appeal_block += f"\n\nТекст апелляционной жалобы:\n{appeal_text_full[:6000]}"
        if extra_instructions.strip():
            appeal_block += f"\n\nДополнительные указания:\n{extra_instructions}"

        doc_rag_query = f"ст. 330 ГПК РФ апелляционное определение {doc_type_tab} {appeal_text_full[:300]}"
        doc_rag = ""
        try:
            doc_rag = get_hybrid_context(doc_rag_query, top_k=st.session_state.get("top_k", 7))
        except Exception:
            pass
        rag_block = f"\n<RAG_CONTEXT>\n{doc_rag}\n</RAG_CONTEXT>\n" if doc_rag else ""

        doc_cache_key = _make_cache_key(
            appeal_block + judge_logic_text, decision_text_full, doc_rag, 0,
        )

        if doc_cache_key in LLM_RESPONSE_CACHE:
            doc_answer = LLM_RESPONSE_CACHE[doc_cache_key]
            st.info("⚡ Документ загружен из кэша")
        else:
            ap_args = appeal_block.strip() or "[Доводы жалобы не предоставлены — проведи анализ самостоятельно]"
            if is_otmena:
                user_doc_message = SYSTEM_PROMPT_OTMENA.format(
                    appeal_arguments=ap_args,
                    rag_context_block=rag_block,
                    source_decision=decision_text_full[:12000],
                )
            else:
                user_doc_message = SYSTEM_PROMPT_BI.format(
                    appeal_arguments=ap_args,
                    judge_logic=judge_logic_text.strip() or "[Не указаны — обоснуй самостоятельно]",
                    rag_context_block=rag_block,
                    source_decision=decision_text_full[:12000],
                )

            with st.spinner("⏳ Генерируем проект определения..."):
                doc_answer = call_llm_streaming(
                    user_doc_message, [],
                    temperature=temp_doc,
                    max_tokens=8000,
                )
            LLM_RESPONSE_CACHE[doc_cache_key] = doc_answer

        st.session_state["generated_doc"] = doc_answer
        st.rerun()

    if st.session_state.get("generated_doc"):
        st.markdown("---")
        st.subheader("📄 Проект апелляционного определения")
        st.markdown(st.session_state["generated_doc"])
        st.divider()

        doc_bytes = create_docx_from_text(st.session_state["generated_doc"])
        fname = "proekt_otmena.docx" if is_otmena else "proekt_bi.docx"
        col_dl, col_clr = st.columns(2)
        with col_dl:
            st.download_button(
                "⬇️ Скачать .docx",
                data=doc_bytes,
                file_name=fname,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        with col_clr:
            if st.button("🗑️ Очистить проект", key="clear_doc", use_container_width=True):
                st.session_state.pop("generated_doc", None)
                st.session_state.pop("draft_request", None)
                st.rerun()

# ══════════════════════════════════════════════════════════════════════════
# ВКЛАДКА 3: БАЗА ЗНАНИЙ
# ══════════════════════════════════════════════════════════════════════════
with tab_db:
    st.subheader("📚 База нормативных документов")
    st.info("ℹ️ Для управления базой используйте вкладку **🛠️ Админ**: импорт документов, допиндексация, пересборка, backup.zip.")
    st.markdown("**Состояние базы:**")
    try:
        kb_info = get_kb_stats()
        st.json(kb_info)
    except Exception:
        st.info("База пуста или статистика недоступна.")

with tab_admin:
    render_admin_tab()
